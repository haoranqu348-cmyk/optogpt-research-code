from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path("/Users/quhaoran/lab/optogpt--")
OUT = ROOT / "joint_284d_paper_manuscript_with_figures.docx"
FIGS = ROOT / "paper_figures/rendered"
NAVY, BLUE, GRAY, GREEN, LIGHT = "17324D", "2E74B5", "666D75", "2C8C62", "E8EEF5"

def set_font(run, size=None, bold=None, color=None):
    run.font.name = "Arial"
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.rFonts
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for key in ("ascii", "hAnsi", "eastAsia"):
        rf.set(qn("w:" + key), "Arial")

def setup(doc):
    sec = doc.sections[0]
    sec.top_margin, sec.bottom_margin = Inches(.8), Inches(.75)
    sec.left_margin, sec.right_margin = Inches(.9), Inches(.9)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Arial", Pt(10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        st = doc.styles[name]
        st.font.name, st.font.size, st.font.bold = "Arial", Pt(size), True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before, st.paragraph_format.space_after = Pt(before), Pt(after)
        st.paragraph_format.keep_with_next = True
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Joint 284D s+p OptoGPT | Manuscript draft"), 8, color=GRAY)

def p(doc, text="", style=None, align=None):
    para = doc.add_paragraph(style=style)
    if align is not None: para.alignment = align
    set_font(para.add_run(text))
    return para

def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(3)
        set_font(para.add_run(item))

def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tcpr.append(shd)
    shd.set(qn("w:fill"), fill)

def cell_margins(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar"); tcpr.append(tcmar)
    for name, value in (("top", 90), ("start", 120), ("bottom", 90), ("end", 120)):
        node = tcmar.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name); tcmar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style, table.alignment = "Table Grid", WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = value
        shade(cell, LIGHT); cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs: set_font(run, 9, True)
    trpr = table.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader"); header.set(qn("w:val"), "true"); trpr.append(header)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value); cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cells[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs: set_font(run, 9)
    p(doc, "")
    return table

def add_figure(doc, filename, label, caption):
    para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before, para.paragraph_format.space_after = Pt(8), Pt(4)
    inline = para.add_run().add_picture(str(FIGS / filename), width=Inches(6.65))
    inline._inline.docPr.set("title", label)
    inline._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_after = Pt(9)
    set_font(cap.add_run(label + "  "), 9.5, True, NAVY)
    set_font(cap.add_run(caption), 9.5)

def ref(doc, n, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(.25)
    para.paragraph_format.first_line_indent = Inches(-.25)
    para.paragraph_format.space_after = Pt(3)
    set_font(para.add_run(f"[{n}] {text}"), 9)

doc = Document(); setup(doc)

# Front matter
para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER; para.paragraph_format.space_before = Pt(42)
set_font(para.add_run("Physics-Validated Generative Inverse Design"), 20, True, NAVY)
para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(para.add_run("An OptoGPT Extension for Dielectric Multilayer Films under Joint s- and p-Polarization Constraints"), 14, color=BLUE)
para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER; para.paragraph_format.space_before = Pt(18)
set_font(para.add_run("Author(s): __________________    Affiliation: __________________"), 11)
para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(para.add_run("Manuscript draft | August 2026"), 10, color=GRAY)
p(doc, "This manuscript consistently describes a single shared-structure joint 284D condition [Rs, Ts, Rp, Tp]. Figures 2 and 3 are presented as training and validation results of the same joint model.", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

p(doc, "Abstract", "Heading 1")
p(doc, "Inverse design of optical multilayer films requires searching a high-dimensional discrete space of materials, layer counts, and thicknesses. Under oblique incidence, the Fresnel response and interference behavior of s- and p-polarized light differ substantially, making independent polarization-specific predictions insufficient when one physical structure must satisfy both channels. Here we extend the OptoGPT generative inverse-design framework to a joint s+p model for dielectric multilayer films at 60° incidence. The model is conditioned on a 284-dimensional optical record [Rs, Ts, Rp, Tp] sampled from 400–1100 nm at 10 nm intervals. Two polarization-aware spectrum branches are fused to drive a shared autoregressive Transformer decoder that emits one material–thickness token sequence. A dielectric logits mask enforces the material contract during inference, while multiple candidates are independently recomputed with the transfer-matrix method (TMM) and ranked by their joint optical error. We construct 500,000 structure-level records with deterministic hashing and zero detected split leakage. On an independent held-out sample, the model achieves 100% valid decoding and 100% TMM verification success, with a mean joint spectral error of 0.03615. Strict out-of-distribution tests yield mean joint MAEs of 0.0930, 0.0719, and 0.0634 for one, 16, and 64 candidates, respectively. A finite-glass high-transmission study shows that p polarization is easier to match than s polarization, while angle sweeps reveal substantial degradation beyond the 60° training condition. The resulting workflow provides a reproducible, physics-validated baseline for joint-polarization thin-film inverse design and makes its current limitations explicit.")
p(doc, "Keywords: multilayer thin films; inverse design; OptoGPT; joint s/p polarization; transfer-matrix method; physics validation; out-of-distribution generalization")

p(doc, "1 Introduction", "Heading 1")
p(doc, "Multilayer dielectric films are used in antireflection coatings, spectral filters, distributed Bragg reflectors, imaging optics, and photovoltaic devices. Given a desired reflectance or transmittance spectrum, a designer must determine both the material ordering and the thickness of every layer. Relative to forward simulation, this inverse problem is discrete, non-convex, high-dimensional, and non-unique: distinct layer sequences can produce nearly indistinguishable optical responses, while local optimization can depend strongly on initialization and on the selected material set.")
p(doc, "The transfer-matrix method (TMM) provides an efficient and accurate forward solver for planar stratified media, but target-by-target TMM optimization still requires repeated simulation calls and may return only one local solution. Deep-learning inverse design reduces the cost of proposing structures by learning a mapping from optical targets to serialized layer sequences. OptoGPT represents a material–thickness pair as a structure token and treats variable-length multilayer design as autoregressive sequence generation, enabling flexible material combinations, layer counts, and stochastic candidate sampling [1]. Conditional invertible neural networks further demonstrate that explicitly representing the multiplicity of inverse solutions can produce useful ensembles rather than a single local optimum [2].")
p(doc, "Joint polarization design introduces an additional physical constraint. At oblique incidence, the s and p channels experience different interface coefficients, phase accumulation, and angular sensitivity. Predicting one structure for each polarization therefore does not guarantee a physically consistent shared film. A valid joint dataset must pair the two responses generated from the same structure, and a valid inference pipeline must recompute both responses with TMM rather than treating neural-network probability as an optical guarantee.")
p(doc, "In this work we develop a unified joint 284D s+p OptoGPT model. Our contributions are:")
bullets(doc, [
    "A structure-level joint data contract that concatenates [Rs, Ts, Rp, Tp] along the feature axis and uses hashing for deduplication and deterministic splitting.",
    "A dual-branch polarization encoder and fusion module connected to a shared autoregressive structure decoder.",
    "A dielectric logits mask and exact s+p TMM reranking for physically valid candidate generation.",
    "An evaluation protocol that separates in-distribution accuracy, strict OOD behavior, candidate-budget gains, structural non-uniqueness, and finite-glass application limits.",
])

p(doc, "2 Physical Problem and Joint Dataset", "Heading 1")
p(doc, "2.1 Optical system and spectral representation", "Heading 2")
p(doc, "We consider an air–dielectric-multilayer–glass system illuminated at a fixed incidence angle of 60°. The wavelength range is 400–1100 nm, sampled every 10 nm at 71 points. The allowed dielectric vocabulary contains Al2O3, AlN, HfO2, MgF2, MgO, Si3N4, SiO2, Ta2O5, TiO2, and ZnO. A structure is serialized as a sequence of material–thickness tokens. For each structure, s- and p-polarized TMM calculations produce Rs, Ts, Rp, and Tp. The four 71-point spectra are concatenated along the feature axis to form one shared 284D condition:")
p(doc, "[Rs(λ1), …, Rs(λ71), Ts(λ1), …, Ts(λ71), Rp(λ1), …, Rp(λ71), Tp(λ1), …, Tp(λ71)] ∈ R^284.", align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, "Figure 1 summarizes the physical application context, the polarization split at oblique incidence, and the proposed generation-and-validation workflow. Every target produces one shared structure; its two polarization responses are recomputed independently by TMM.")
add_figure(doc, "figure1.png", "Figure 1", "Physical problem, joint s+p OptoGPT workflow, and spectral comparisons between classical and AI-assisted coating designs. The proposed workflow includes polarization-aware conditioning, dielectric candidate generation, and exact TMM reranking.")

p(doc, "2.2 Data generation, deduplication, and splitting", "Heading 2")
p(doc, "Legal material–thickness sequences are sampled first, after which the same physical structure is simulated twice, once for each polarization. The four spectra are concatenated only along the feature axis; s and p samples are never mixed along the sample axis. Each record stores the token sequence, layer metadata, incidence angle, wavelength grid, and joint spectrum. A structure-level SHA-256 hash is used for deduplication and deterministic splitting, so a physical structure and its identical optical record cannot appear in more than one split.")
add_table(doc, ["Subset", "Structures", "Spectrum dimension", "Structure leakage"], [
    ["Training", "400,006", "284", "0"],
    ["Development", "50,104", "284", "0"],
    ["Test", "49,890", "284", "0"],
    ["Total", "500,000", "284", "0"],
])
p(doc, "2.3 Evaluation metrics", "Heading 2")
p(doc, "For a target spectrum y and a TMM-recomputed candidate spectrum ŷ, we calculate polarization-channel mean absolute errors Es and Ep and define the joint error as Ejoint=(Es+Ep)/2. We additionally record valid decoding rate, TMM success rate, mean transmission, p05 transmission, and the number of TMM calls associated with each candidate budget. MAE is a spectral matching metric; it is not a structure-recovery accuracy.")

p(doc, "3 Joint 284D OptoGPT Method", "Heading 1")
p(doc, "3.1 Structure serialization and shared decoder", "Heading 2")
p(doc, "Each layer is represented by one material_thickness token. Tokens are ordered according to the layer sequence and bounded by BOS and EOS markers. A shared decoder-only Transformer generates the sequence autoregressively from the fused optical condition. This representation removes the fixed-output-size limitation of conventional regressors and permits different material orderings, layer counts, and thickness combinations.")
p(doc, "3.2 Dual polarization branches and parameter transfer", "Heading 2")
p(doc, "The 284D condition is partitioned into an s branch [Rs, Ts] and a p branch [Rp, Tp]. Separate spectrum encoders extract channel-aware representations, which are then fused into a shared condition embedding. The token embedding, six-layer autoregressive decoder, and output generator are transferred from pretrained OptoGPT, while the new polarization branches and fusion parameters are initialized for the joint task. Training uses a fusion warm-up stage followed by full fine-tuning, reducing disruption to the inherited structure representation.")
p(doc, "3.3 Dielectric constraints and physics reranking", "Heading 2")
p(doc, "During inference, logits for metals and semiconductors are masked to negative infinity, enforcing the ten-material dielectric contract. For each target, N candidates are sampled, checked for sequence legality, deduplicated, and independently recomputed with s/p TMM. Candidates are ranked by Ejoint. The neural model therefore acts as a fast proposal mechanism, whereas TMM remains the physical evaluator.")
p(doc, "3.4 Training configuration and reproducibility", "Heading 2")
p(doc, "Formal training uses 500,000 joint samples, batch size 16, Adam, mixed precision, label smoothing, ReduceLROnPlateau scheduling, and differential learning rates. The best checkpoint is selected by development loss and stores model, optimizer, scheduler, AMP scaler, epoch and step, random-number states, architecture version, pretrained-weight hash, and the data manifest. Figure 2 presents the joint training progression, dielectric vocabulary constraint, and end-to-end physics-evaluation contract.")
add_figure(doc, "figure2.png", "Figure 2", "Training progression, channel-wise performance, dielectric-vocabulary control, and the joint physics-evaluation contract of the 284D s+p model. All four performance bars belong to the same shared-structure joint model.")

p(doc, "4 Training Convergence and Joint-Model Validation", "Heading 1")
p(doc, "4.1 Training convergence", "Heading 2")
p(doc, "The joint model train loss decreases from 4.6647 to 3.3164, while development loss decreases from 4.3498 to 3.0381. Both phases show continued improvement, with the best development checkpoint at epoch 10. Token-prediction loss should be distinguished from optical MAE: the former measures sequence-model optimization, whereas the latter is obtained only after generating structures and recomputing their spectra with TMM.")
p(doc, "4.2 Spectral reconstruction and structural statistics", "Heading 2")
p(doc, "Figure 3 shows best and typical validation reconstructions from the joint model. Gray curves represent the target optical responses, and colored dashed curves represent the s/p channel responses recomputed from the generated shared structures. Across 200 valid joint-model samples, the mean total MAE is 0.03157 with a standard deviation of 0.02157. The retained structures use the ten allowed dielectric materials, contain 1–15 layers, and have a mean generated layer count of 6.50. TiO2, ZnO, Ta2O5, and Si3N4 are among the most frequently used materials, indicating non-uniform preferences within the discrete vocabulary.")
add_figure(doc, "figure3.png", "Figure 3", "Joint 284D spectral reconstruction, total-MAE distribution, material usage, layer-count statistics, and wavelength-resolved s/p errors. Every panel refers to the same shared-structure joint model.")
p(doc, "4.3 Independent held-out performance", "Heading 2")
add_table(doc, ["Metric", "Result", "Interpretation"], [
    ["Valid decoding", "100/100", "All sampled targets produced legal structures"],
    ["TMM validation", "100/100", "All candidates completed exact s/p recomputation"],
    ["Es", "0.05084", "Mean s-channel spectral error"],
    ["Ep", "0.02145", "Mean p-channel spectral error"],
    ["Ejoint", "0.03615", "Average of the two channel errors"],
    ["Mean Ts", "0.50367", "Mean s-channel transmission over mixed targets"],
    ["Mean Tp", "0.90179", "Mean p-channel transmission over mixed targets"],
])
p(doc, "The random held-out pool contains diverse target spectra and is not a dedicated flat high-transmission benchmark. Mean Ts should therefore not be interpreted as an upper bound on antireflection performance; it primarily indicates that the s channel is the larger source of joint error.")

p(doc, "5 Candidate Budgets, OOD Generalization, and Non-Uniqueness", "Heading 1")
p(doc, "5.1 Candidate-budget gains", "Heading 2")
p(doc, "To separate one-shot generation from physics-guided search, we compare greedy decoding with 16-candidate and 64-candidate sampling. Every candidate follows the same legality, deduplication, and TMM evaluation procedure, and the best candidate error is reported for each target.")
add_table(doc, ["Strategy", "Mean joint MAE", "Median MAE", "Worst MAE"], [
    ["1 candidate, greedy", "0.0930", "0.0889", "0.1722"],
    ["16 candidates + TMM ranking", "0.0719", "0.0714", "0.1300"],
    ["64 candidates + TMM ranking", "0.0634", "0.0643", "0.1131"],
])
p(doc, "Increasing the candidate budget from one to 64 reduces mean joint MAE from 0.0930 to 0.0634. This improvement is a system-level effect of stochastic structural diversity and exact physical selection, rather than a property of the Transformer alone.")
p(doc, "5.2 Strict out-of-distribution testing", "Heading 2")
p(doc, "The strict OOD set contains continuous rather than 10 nm-quantized thicknesses, 15–20-layer structures, random, graded, strongly alternating, and bimodal thickness patterns, and small spectral perturbations for a subset of targets. Hash checks against the verifiable training structures prevent exact duplicates. With 64 candidates, the mean s-channel MAE is 0.0827 and the mean p-channel MAE is 0.0440. Strongly alternating structures are the most difficult, indicating sensitivity to rapid phase variation and complex sequences outside the training distribution.")
p(doc, "5.3 Spectral equivalence versus structural recovery", "Heading 2")
p(doc, "The exact original structure is recovered for 0 of 60 strict OOD targets. This does not imply that inverse design has failed: different material orderings and thicknesses can produce nearly equivalent optical responses. The result instead shows that the model primarily finds a physically equivalent design, not the hidden structure used to generate the target. This distinction is central for interpreting generative inverse design results.")

p(doc, "6 Finite-Glass High-Transmission Design and Capability Boundaries", "Heading 1")
p(doc, "To assess a more device-like substrate condition, candidates are evaluated on a 500 μm finite glass substrate using a finite-substrate TMM contract that includes incoherent internal substrate reflections. The reference target is flat over 400–1100 nm with Rs=Rp=0.05 and Ts=Tp=0.95.")
add_figure(doc, "figure4.png", "Figure 4", "Three joint s+p high-transmission candidates on finite glass. All candidates are driven by the same flat reference target and evaluated with finite-glass TMM. The best retained candidate reaches mean Ts=0.773 and mean Tp=0.982 and does not attain the common 0.95 target.")
p(doc, "Figure 4 shows that the p channel can approach the reference more closely than the s channel, while the latter retains substantial reflection loss under the finite-glass contract. This result should be interpreted as an engineering capability boundary, not as evidence that broadband dual-polarization 95% transmission has been solved. Finite-substrate reflections, the glass–coating interfaces, and angular polarization asymmetry jointly increase the design difficulty.")
p(doc, "Angle sweeps further show degradation beyond the training condition. For a representative candidate, mean Ts is approximately 0.3525 and mean Tp approximately 0.5925 near 80°. The model should therefore be described as a 60° joint model; its single-angle result must not be generalized to 0–75° wide-angle performance.")

p(doc, "7 Discussion", "Heading 1")
p(doc, "The main strength of the present framework is that joint polarization is formulated as a shared-structure generation problem with an explicit data contract and a physics-based evaluation loop. Unlike reporting neural-network loss alone, the workflow exposes legality, TMM success, candidate budget, polarization asymmetry, and OOD degradation to independent checks. Three limitations are particularly important:")
bullets(doc, [
    "The s channel is the dominant bottleneck at 60° and higher angles, suggesting that the current dense-dielectric vocabulary and discrete thickness space may not contain sufficiently broad high-angle antireflection solutions.",
    "Multiple candidates and TMM reranking improve performance but increase simulation cost; future comparisons should normalize results by TMM-call budget.",
    "The model generates spectrally equivalent structures rather than a unique recovered ground-truth film. Manufacturing-oriented design requires additional minimum-thickness, total-thickness, material-compatibility, and tolerance constraints.",
])
p(doc, "Promising next steps include multi-angle conditioning, graded or porous refractive-index materials, manufacturing-error loops, and active-learning acquisition. The current active-learning implementation provides validated records and deduplication, but ensemble uncertainty, acquisition, closed-loop retraining, and controlled comparisons are not yet complete and are not claimed here.")

p(doc, "8 Conclusion", "Heading 1")
p(doc, "We presented and evaluated a joint s+p OptoGPT model for dielectric multilayer inverse design at 60° incidence. The model conditions on a shared 284D [Rs, Ts, Rp, Tp] spectrum, uses polarization-aware encoding and shared autoregressive decoding, enforces the dielectric vocabulary during inference, and reranks generated candidates with exact s+p TMM. The 500,000-record structure-level dataset, 100% sampled decoding and TMM success rates, and the monotonic improvement obtained with larger candidate budgets establish a reproducible evidence chain for the method.")
p(doc, "The results also make the current boundary clear: s polarization remains the main source of error, the finite-glass 95% dual-polarization target is not fully reached, and a model trained at 60° degrades at larger angles. By reporting both the capability and these limitations, this work provides a physics-checkable baseline for future wide-angle, materials-expanded, and active-learning thin-film inverse-design studies.")

p(doc, "References", "Heading 1")
references = [
    "Ma T, Wang H, Guo LJ. OptoGPT: A foundation model for inverse design in optical multilayer thin film structures. Opto-Electronic Advances. 2024;7(7):240062. doi:10.29026/oea.2024.240062.",
    "Luce A, Mahdavi A, Wankerl H, Marquardt F. Investigation of inverse design of multilayer thin-films with conditional invertible neural networks. Machine Learning: Science and Technology. 2023;4(1):015014. doi:10.1088/2632-2153/acb48d.",
    "Hong Y, Nicholls DP. Data-driven design of thin-film optical systems using deep active learning. Optics Express. 2022;30(13):22901. doi:10.1364/OE.459295.",
    "Khaireh-Walieh A, Langevin D, Bennet P, Teytaud O, Moreau A, Wiecha PR. A newcomer’s guide to deep learning for inverse design in nano-photonics. Nanophotonics. 2023;12(24):4387–4414. doi:10.1515/nanoph-2023-0527.",
    "Wiecha PR, Arbouet A, Girard C, Muskens OL. Deep learning in nano-photonics: inverse design and beyond. Photonics Research. 2021;9(5):B182–B200. doi:10.1364/PRJ.415960.",
    "Byrnes SJ. Multilayer optical calculations. arXiv:1603.02720. 2016. doi:10.48550/arXiv.1603.02720.",
    "Macleod HA. Thin-Film Optical Filters. 5th ed. Boca Raton: CRC Press; 2017.",
    "Goodfellow I, Bengio Y, Courville A. Deep Learning. Cambridge: MIT Press; 2016.",
    "Settles B. Active Learning Literature Survey. Madison: University of Wisconsin–Madison; 2009.",
]
for i, item in enumerate(references, 1): ref(doc, i, item)

doc.save(OUT)
print(OUT)

