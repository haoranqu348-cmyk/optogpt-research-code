from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path("/Users/quhaoran/lab/optogpt--")
OUT = ROOT / "Physics_Validated_JP_OptoGPT_Manuscript_Body.docx"
FIGS = ROOT / "paper_figures" / "rendered"
NAVY, BLUE, GRAY, LIGHT = "17324D", "2E74B5", "666D75", "E8EEF5"

def set_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Arial"
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn("w:" + key), "Arial")

def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.2)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5.5)
    for name, size, color, before, after in [
        ("Heading 1", 15, BLUE, 14, 7),
        ("Heading 2", 12.5, BLUE, 10, 5),
        ("Heading 3", 11, NAVY, 7, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Physics-Validated JP-OptoGPT | Manuscript body"), 8, color=GRAY)
    settings = doc.settings._element
    math_pr = settings.find(qn("m:mathPr"))
    if math_pr is None:
        math_pr = OxmlElement("m:mathPr")
        settings.append(math_pr)
    math_font = math_pr.find(qn("m:mathFont"))
    if math_font is None:
        math_font = OxmlElement("m:mathFont")
        math_pr.append(math_font)
    math_font.set(qn("m:val"), "Cambria Math")

def add_paragraph(doc, text="", style=None, align=None):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    set_font(paragraph.add_run(text))
    return paragraph

def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        set_font(paragraph.add_run(item))

def math_run(text, style=None):
    run = OxmlElement("m:r")
    if style is not None:
        math_rpr = OxmlElement("m:rPr")
        math_style = OxmlElement("m:sty")
        math_style.set(qn("m:val"), style)
        math_rpr.append(math_style)
        run.append(math_rpr)
    word_rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn("w:" + key), "Cambria Math")
    word_rpr.append(fonts)
    run.append(word_rpr)
    text_node = OxmlElement("m:t")
    text_node.text = text
    run.append(text_node)
    return run

def math_nodes(*items):
    nodes = []
    for item in items:
        if item is None:
            continue
        if isinstance(item, (list, tuple)):
            nodes.extend(math_nodes(*item))
        elif isinstance(item, str):
            nodes.append(math_run(item))
        else:
            nodes.append(item)
    return nodes

def math_slot(name, content):
    slot = OxmlElement("m:" + name)
    for node in math_nodes(content):
        slot.append(node)
    return slot

def math_sub(base, subscript):
    node = OxmlElement("m:sSub")
    node.append(OxmlElement("m:sSubPr"))
    node.append(math_slot("e", base))
    node.append(math_slot("sub", subscript))
    return node

def math_sup(base, superscript):
    node = OxmlElement("m:sSup")
    node.append(OxmlElement("m:sSupPr"))
    node.append(math_slot("e", base))
    node.append(math_slot("sup", superscript))
    return node

def math_fraction(numerator, denominator):
    node = OxmlElement("m:f")
    properties = OxmlElement("m:fPr")
    fraction_type = OxmlElement("m:type")
    fraction_type.set(qn("m:val"), "bar")
    properties.append(fraction_type)
    node.append(properties)
    node.append(math_slot("num", numerator))
    node.append(math_slot("den", denominator))
    return node

def math_delimiter(content, begin="(", end=")"):
    node = OxmlElement("m:d")
    properties = OxmlElement("m:dPr")
    begin_char = OxmlElement("m:begChr")
    begin_char.set(qn("m:val"), begin)
    end_char = OxmlElement("m:endChr")
    end_char.set(qn("m:val"), end)
    properties.extend([begin_char, end_char])
    node.append(properties)
    node.append(math_slot("e", content))
    return node

def math_matrix(rows):
    matrix = OxmlElement("m:m")
    matrix.append(OxmlElement("m:mPr"))
    for row in rows:
        matrix_row = OxmlElement("m:mr")
        for cell in row:
            matrix_row.append(math_slot("e", cell))
        matrix.append(matrix_row)
    return matrix

def math_sum(expression, lower, upper):
    node = OxmlElement("m:nary")
    properties = OxmlElement("m:naryPr")
    char = OxmlElement("m:chr")
    char.set(qn("m:val"), "∑")
    limit_location = OxmlElement("m:limLoc")
    limit_location.set(qn("m:val"), "undOvr")
    properties.extend([char, limit_location])
    node.append(properties)
    node.append(math_slot("sub", lower))
    node.append(math_slot("sup", upper))
    node.append(math_slot("e", expression))
    return node

def math_bar(content):
    node = OxmlElement("m:bar")
    properties = OxmlElement("m:barPr")
    position = OxmlElement("m:pos")
    position.set(qn("m:val"), "top")
    properties.append(position)
    node.append(properties)
    node.append(math_slot("e", content))
    return node

def math_function(name, argument):
    node = OxmlElement("m:func")
    node.append(OxmlElement("m:funcPr"))
    function_name = math_run(name, "p") if isinstance(name, str) else name
    node.append(math_slot("fName", function_name))
    node.append(math_slot("e", argument))
    return node

def math_apply(name, argument):
    return math_nodes(name, math_delimiter(argument))

def symbol(name, subscript=None, superscript=None, style=None):
    node = math_run(name, style)
    if subscript is not None:
        node = math_sub(node, subscript)
    if superscript is not None:
        node = math_sup(node, superscript)
    return node

def equation_formula(number):
    lam = lambda sub=None: symbol("λ", sub)
    theta = lambda sub=None: symbol("θ", sub)
    n = lambda sub=None: symbol("n", sub)
    q = lambda sub=None: symbol("q", sub)
    d = lambda sub=None: symbol("d", sub)
    delta = lambda sub=None: symbol("δ", sub)
    if number == 1:
        spectrum = lambda name: math_apply(
            symbol(name[0], name[1]),
            lam(math_nodes("1", math_run(":"), symbol("N", "λ"))),
        )
        return math_nodes(
            math_apply(math_run("c", "b"), "x"), " = ",
            math_delimiter([spectrum("Rs"), ", ", spectrum("Ts"), ", ", spectrum("Rp"), ", ", spectrum("Tp")], "[", "]"),
            " ∈ ", math_sup(math_run("ℝ", "p"), math_nodes("4", symbol("N", "λ"))),
            " = ", math_sup(math_run("ℝ", "p"), "284"),
        )
    if number == 2:
        return math_nodes(n("0"), " ", math_function("sin", theta("0")), " = ", n("j"), " ", math_function("sin", theta("j")))
    if number == 3:
        return math_nodes(
            math_apply(delta("j"), "λ"), " = ", math_fraction(math_nodes("2π"), "λ"), " ",
            math_apply(n("j"), "λ"), " ", d("j"), " ", math_function("cos", theta("j")),
        )
    if number == 4:
        return math_nodes(
            q(math_nodes("j", math_run(","), math_run("s", "p"))), " = ", n("j"), " ", math_function("cos", theta("j")),
            ",      ", q(math_nodes("j", math_run(","), math_run("p", "p"))), " = ",
            math_fraction(n("j"), math_function("cos", theta("j"))),
        )
    if number == 5:
        matrix = math_matrix([
            [math_function("cos", delta("j")), math_fraction(math_nodes("i ", math_function("sin", delta("j"))), q("j"))],
            [math_nodes("i ", q("j"), " ", math_function("sin", delta("j"))), math_function("cos", delta("j"))],
        ])
        return math_nodes(symbol("M", "j"), " = ", math_delimiter(matrix, "[", "]"))
    if number == 6:
        return math_nodes(
            "r = ", math_fraction(math_nodes(q("0"), "B − C"), math_nodes(q("0"), "B + C")),
            ",      t = ", math_fraction(math_nodes("2", q("0")), math_nodes(q("0"), "B + C")),
        )
    if number == 7:
        absolute = lambda value: math_delimiter(value, "|", "|")
        return math_nodes(
            "R = ", math_sup(absolute("r"), "2"), ",      T = ",
            math_fraction(math_apply(math_run("Re", "p"), q(math_run("sub", "p"))), math_apply(math_run("Re", "p"), q("0"))),
            " ", math_sup(absolute("t"), "2"),
        )
    if number == 8:
        channel = lambda name, suffix: math_apply(
            symbol(name, math_nodes("q", math_run(","), math_run(suffix, "p"))), lam("k")
        )
        residuals = math_nodes(
            math_delimiter(math_nodes(channel("R", "target"), " − ", channel("R", "x")), "|", "|"),
            " + ",
            math_delimiter(math_nodes(channel("T", "target"), " − ", channel("T", "x")), "|", "|"),
        )
        return math_nodes(
            symbol("E", "q"), " = ", math_fraction("1", math_nodes("2", symbol("N", "λ"))), " ",
            math_sum(math_delimiter(residuals), math_nodes("k = 1"), symbol("N", "λ")),
        )
    if number == 9:
        return math_nodes(symbol("E", math_run("joint", "p")), " = ", math_fraction(math_nodes(symbol("E", "s"), " + ", symbol("E", "p")), "2"))
    if number == 10:
        minimizer = math_sub(math_run("arg min", "p"), math_nodes(symbol("x", "i"), " ∈ ", symbol("C", math_run("legal", "p"))))
        return math_nodes(math_sup("x", "*"), " = ", minimizer, " ", math_apply(symbol("E", math_run("joint", "p")), symbol("x", "i")))
    if number == 11:
        mean_t = lambda sub: math_sub(math_bar("T"), sub)
        summed = math_sum(math_apply(symbol("T", "q"), lam("k")), "k = 1", symbol("N", "λ"))
        return math_nodes(
            mean_t("q"), " = ", math_fraction("1", symbol("N", "λ")), " ", summed,
            ",      ", symbol("T", math_run("worst", "p")), " = ",
            math_apply(math_run("min", "p"), math_nodes(mean_t("s"), ", ", mean_t("p"))),
        )
    if number == 12:
        mean_t = lambda sub: math_sub(math_bar("T"), sub)
        return math_nodes(
            mean_t(math_run("unpol", "p")), " = ",
            math_fraction(math_nodes(mean_t("s"), " + ", mean_t("p")), "2"),
        )
    raise ValueError("Unsupported equation number: " + str(number))

def add_equation(doc, number):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(3.35), WD_TAB_ALIGNMENT.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.55), WD_TAB_ALIGNMENT.RIGHT)
    set_font(paragraph.add_run("\t"), 10.5)
    equation = OxmlElement("m:oMath")
    for node in equation_formula(number):
        equation.append(node)
    paragraph._p.append(equation)
    set_font(paragraph.add_run("\t(" + str(number) + ")"), 10.5)
    return paragraph

def shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shading = tcpr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tcpr.append(shading)
    shading.set(qn("w:fill"), fill)

def set_cell_margins(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    margins = tcpr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)
    for name, value in (("top", 80), ("start", 110), ("bottom", 80), ("end", 110)):
        node = margins.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def add_table(doc, label, caption, headers, rows):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(5)
    cap.paragraph_format.space_after = Pt(4)
    set_font(cap.add_run(label + ". "), 9.5, bold=True, color=NAVY)
    set_font(cap.add_run(caption), 9.5)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = value
        shade_cell(cell, LIGHT)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, 8.7, bold=True)
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    row_properties.append(header_flag)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            set_cell_margins(cells[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_font(run, 8.5)
    add_paragraph(doc, "")
    return table

def add_figure(doc, filename, label, caption):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    inline = paragraph.add_run().add_picture(str(FIGS / filename), width=Inches(6.75))
    inline._inline.docPr.set("title", label)
    inline._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_after = Pt(8)
    set_font(cap.add_run(label + ". "), 9, bold=True, color=NAVY)
    set_font(cap.add_run(caption), 9)

def add_reference(doc, number, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.24)
    paragraph.paragraph_format.first_line_indent = Inches(-0.24)
    paragraph.paragraph_format.space_after = Pt(1.2)
    paragraph.paragraph_format.line_spacing = 1.0
    set_font(paragraph.add_run("[" + str(number) + "] " + text), 8.2)

doc = Document()
setup_document(doc)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(14)
set_font(title.add_run("Physics-Validated Joint-Polarization Generative Inverse Design of Dielectric Multilayer Films"), 18, bold=True, color=NAVY)

add_paragraph(doc, "1. Introduction", style="Heading 1")
add_paragraph(doc, "Optical multilayer films provide one of the most versatile platforms for controlling reflection, transmission, absorption, phase, and color over broad spectral ranges. Their planar geometry is compatible with mature deposition processes, yet the corresponding design space is unusually large because a complete structure requires the simultaneous selection of materials, layer ordering, layer count, and individual thicknesses. Even when the material set is fixed, the number of possible sequences grows combinatorially and the spectral response varies nonlinearly with thickness through multiple reflection and interference. Conventional thin-film synthesis therefore combines analytical intuition with numerical refinement, needle insertion, evolutionary search, or other target-specific optimization procedures [1]. Such methods remain powerful, but a new optimization must usually be performed for every target, and the returned design can depend strongly on the initial condition and the prescribed structural family.")
add_paragraph(doc, "Machine-learning approaches have introduced a complementary strategy in which an inverse model learns reusable mappings from desired optical responses to candidate structures. Reviews of nanophotonic inverse design have shown that discriminative networks, tandem architectures, mixture-density models, generative models, reinforcement learning, and differentiable surrogate solvers address different aspects of this problem [2–4]. For multilayer films in particular, deep reinforcement learning has been used to construct variable-length coatings [5,6], and probabilistic models have been developed to represent one-to-many mappings between spectra and layered structures [7]. These developments reduce the cost of proposing a design after training, but they do not by themselves guarantee that the generated structure obeys a material contract or that its independently simulated spectrum matches the requested target.")
add_paragraph(doc, "OptoGPT introduced a particularly flexible formulation by combining material and thickness into a structure token and casting thin-film inverse design as conditional autoregressive generation [8]. Unlike fixed-dimensional regressors, this representation can generate variable numbers of layers and multiple structural alternatives from the same target. The original framework also demonstrated that sampling and fine-tuning can adapt a general pretrained model to specialized optical tasks. Nevertheless, extending a generative model to an oblique-incidence coating problem requires more than changing the target spectrum. The physical condition must represent both polarization channels of the same structure, forbidden materials must be excluded deterministically, and every candidate must be judged using an exact optical solver rather than sequence probability alone.")
add_paragraph(doc, "Joint s- and p-polarization design is important because an oblique beam does not experience a polarization-independent interface. At 60° incidence, the Fresnel coefficients, optical admittances, and phase accumulated inside each layer differ for the two channels. A structure that provides low p-polarized reflectance can therefore remain strongly reflective for s polarization. Treating the channels as two unrelated inverse-design problems would produce two different films and would not solve the engineering requirement of one coating under an arbitrarily polarized field. A physically meaningful learning record must instead associate a single layer sequence with its complete [Rs, Ts, Rp, Tp] response.")
add_paragraph(doc, "This work develops a physics-validated joint-polarization OptoGPT framework, denoted JP-OptoGPT, for dielectric multilayer design at 60° incidence. The contribution is a system rather than a single architectural substitution. First, the optical condition is expanded to a structure-paired 284-dimensional record containing reflection and transmission for both polarization channels. Second, polarization-aware spectrum branches and a fusion module condition one shared autoregressive decoder. Third, a dielectric logits mask removes metals and semiconductors from the sampling distribution, guaranteeing a legal material vocabulary at inference. Fourth, JP-OptoGPT is used as a multi-candidate proposal model, while exact s/p transfer-matrix calculations provide the final physical ranking. Finally, the generated candidates are transferred to a finite-glass application and used to initialize an independently refined double-sided coating.")
add_paragraph(doc, "The distinction between model generation and physical selection is central to our evaluation. A single decoded sequence measures one-shot generative behavior; the best of 16 or 64 candidates measures the combined capability of stochastic sampling and TMM reranking. We report both quantities rather than assigning the entire gain to the Transformer. We likewise distinguish the formal held-out test from a strict out-of-distribution (OOD) set containing continuous thicknesses, long sequences, alternating patterns, and perturbed spectra. This separates interpolation within the sampled design grammar from generalization to structures that violate several training regularities.")
add_paragraph(doc, "The article is organized around the four principal figures. Figure 1 introduces the physical problem and contrasts the original single-condition workflow with JP-OptoGPT. Figure 2 presents the joint training trajectory, channel asymmetry, material-vocabulary purification, and evaluation contract. Figure 3 provides spectral and structural statistics, while the formal 100-sample held-out evaluation remains the primary quantitative result. Figure 4 demonstrates finite-glass high-transmission design, compares the model-generated solution with bare glass, and shows a physics-refined double-sided extension. Strict OOD tests and 0–80° angle scans are then used to define the present limits of the method. The objective is not to claim universal wide-angle antireflection, but to establish a reproducible generative design workflow whose capabilities and boundaries are independently verifiable.")

add_paragraph(doc, "2. Joint-Polarization Design Problem", style="Heading 1")
add_paragraph(doc, "2.1 Optical configuration and shared condition", style="Heading 2")
add_paragraph(doc, "The basic configuration is an isotropic multilayer deposited between air and a glass substrate. Unless otherwise specified, the external incidence angle is 60°, the wavelength interval is 400–1100 nm, and the grid spacing is 10 nm, giving Nλ=71 samples. The dielectric vocabulary contains Al2O3, AlN, HfO2, MgF2, MgO, Si3N4, SiO2, Ta2O5, TiO2, and ZnO. A generated layer is represented by a material–thickness token. The training grammar permits variable layer counts and discretized thicknesses, whereas the strict OOD evaluation additionally includes continuous thickness values.")
add_paragraph(doc, "For a structure x, the forward solver returns reflection and transmission spectra for the s and p polarizations. The joint condition is formed by concatenating the four spectra along the feature axis, not by mixing independent polarization records along the sample axis:")
add_equation(doc, 1)
add_paragraph(doc, "Equation (1) enforces the physical pairing that defines the task: all four spectra must originate from the same layer sequence. The ordering is fixed throughout data generation, training, deployment, and evaluation. A 142-dimensional record containing only [R,T] remains valid for a single-polarization task but is not a joint-polarization condition. This data-level distinction is essential because a network can train successfully on a numerically well-formed but physically mispaired dataset.")
add_paragraph(doc, "2.2 Transfer-matrix formulation", style="Heading 2")
add_paragraph(doc, "We use the transfer-matrix method as the authoritative forward solver. For a nonmagnetic layer j with refractive index nj, thickness dj, and internal angle θj, Snell’s law relates the angles in the stratified stack:")
add_equation(doc, 2)
add_paragraph(doc, "The phase thickness accumulated during one pass through layer j is")
add_equation(doc, 3)
add_paragraph(doc, "and the normalized optical admittance depends on polarization:")
add_equation(doc, 4)
add_paragraph(doc, "With the convention of Eq. (4), the characteristic matrix of layer j is")
add_equation(doc, 5)
add_paragraph(doc, "The total matrix M is the ordered product of the layer matrices. If q0 and qsub denote the incident-medium and substrate admittances for the selected polarization, the fields at the first boundary can be written as [B,C]^T=M[1,qsub]^T. The amplitude coefficients are then")
add_equation(doc, 6)
add_paragraph(doc, "and the corresponding intensity coefficients are")
add_equation(doc, 7)
add_paragraph(doc, "The same physical layer sequence is evaluated twice using the polarization-specific admittances. For lossless dielectric stacks on a semi-infinite substrate, R+T is unity up to numerical precision. For the finite-glass application, the coating fields are treated coherently while the 500 μm glass substrate is treated as an incoherent layer. The finite-substrate calculation sums the repeated internal intensity transfers and includes the uncoated or coated rear interface. This distinction matters at 60°: a coating that performs well on the front surface alone can lose part of its benefit after the rear-surface penalty is included. The TMM implementation follows standard multilayer conventions and was independently checked against the TMM-Fast formulation [9] and the widely used multilayer calculation treatment of Byrnes [10].")
add_paragraph(doc, "2.3 Error functions and physical ranking", style="Heading 2")
add_paragraph(doc, "For polarization q∈{s,p}, the spectral error combines the reflection and transmission deviations over the wavelength grid:")
add_equation(doc, 8)
add_paragraph(doc, "The joint error gives equal weight to the two polarization channels:")
add_equation(doc, 9)
add_paragraph(doc, "For a candidate set C={x1,…,xK} sampled from the autoregressive decoder, the selected design is")
add_equation(doc, 10)
add_paragraph(doc, "where Clegal contains only sequences that pass syntax, material, thickness, and TMM checks. Equation (10) makes the role of candidate budget explicit. K=1 is a one-shot result, whereas K>1 combines generation with a finite physical search. Sequence log probability is not used as the final optical objective.")
add_paragraph(doc, "For high-transmission design, we also report the mean transmission of each polarization and the limiting polarization metric:")
add_equation(doc, 11)
add_paragraph(doc, "For an unpolarized reference with equal channel weighting,")
add_equation(doc, 12)
add_paragraph(doc, "Equations (11) and (12) are especially useful for interpreting the finite-glass example. Improving Tbar_s while slightly reducing an already high Tbar_p can increase both Tworst and Tbar_unpol, even though the p-polarized value does not improve independently.")
add_figure(doc, "figure1.png", "Figure 1", "Multilayer thin-film optics and the physics-validated joint-polarization inverse-design workflow. (a) A coated objective lens and enlarged oblique-incidence interface illustrate the distinct s- and p-polarized field orientations. (b) JP-OptoGPT receives one shared [Rs,Ts,Rp,Tp] condition, generates dielectric candidates autoregressively, and uses exact 60° s/p TMM for physical reranking. (c,d) A conventional MgF2 coating is shown at normal and 60° incidence to expose polarization splitting. (e) A model-generated high-transmission candidate demonstrates the application target. The original OptoGPT result is used only as a qualitative single-condition reference; its task and condition dimensionality are not treated as a numerically equivalent baseline.")

add_paragraph(doc, "3. Physics-Validated OptoGPT Extension", style="Heading 1")
add_paragraph(doc, "3.1 Structure serialization and inherited decoder", style="Heading 2")
add_paragraph(doc, "JP-OptoGPT retains the central insight of OptoGPT: a multilayer can be represented as a language-like sequence whose tokens jointly encode material identity and thickness. A structure with L layers is serialized as [BOS, m1_d1, …, mL_dL, EOS]. The decoder learns conditional token distributions and can terminate after different numbers of layers. Compared with direct thickness regression for a fixed material sequence, serialization supports variable layer counts and material choices without changing the output dimension.")
add_paragraph(doc, "The inherited components include the token embedding, positional representation, six-layer autoregressive decoder, and output generator. This transfer is useful because the pretrained weights already encode statistical regularities of layer ordering and thickness combinations. JP-OptoGPT does not claim a new Transformer primitive; its methodological contribution is the joint optical conditioning and the physics-validated inference system built around the transferred structural generator.")
add_paragraph(doc, "3.2 Polarization-aware conditioning", style="Heading 2")
add_paragraph(doc, "The 284D condition is partitioned into two 142D branches. The s branch processes [Rs,Ts], and the p branch processes [Rp,Tp]. Each branch maps its spectrum to a channel representation, after which a fusion module produces the condition supplied to the shared decoder. This design preserves polarization identity before fusion and avoids asking one undifferentiated fully connected layer to infer the segmentation of four spectra. The fused representation conditions one structure sequence, not one sequence per branch.")
add_paragraph(doc, "Training begins with a fusion warm-up phase in which inherited structural parameters are frozen and the newly initialized condition modules adapt to the pretrained decoder. Full fine-tuning then unlocks all parameters, with a higher learning rate assigned to the new modules and a lower rate applied to inherited weights. This two-stage procedure reduces the risk that the new optical condition immediately disrupts the pretrained sequence representation.")
add_paragraph(doc, "3.3 Deterministic dielectric constraint", style="Heading 2")
add_paragraph(doc, "The application requires transparent dielectric coatings. The source vocabulary contains materials that are undesirable for this task, including metals and semiconductors with appreciable extinction coefficients. Material auditing confirmed that Ag, Al, Ge, Si, TiN, and related candidates can introduce strong absorption or metallic reflection over the target band. At every material-generation step, JP-OptoGPT applies a logits mask that sets forbidden material tokens to negative infinity before softmax. The remaining distribution is renormalized over the ten approved dielectrics.")
add_paragraph(doc, "This constraint should be interpreted as a deterministic validity mechanism, not as an empirically demonstrated MAE ablation. The migrated archive does not contain a controlled joint-284D experiment in which the same checkpoint and candidate budget are evaluated with and without the mask. We therefore claim that the mask guarantees material legality, which is directly verifiable, but do not claim a numerical accuracy gain caused by the mask. This distinction avoids conflating an engineering contract with a predictive-performance result.")
add_paragraph(doc, "3.4 Candidate generation and exact reranking", style="Heading 2")
add_paragraph(doc, "At inference, the joint condition is encoded once and the decoder generates K sequences using greedy or stochastic sampling. Candidate validity is checked before optical evaluation. Special tokens are removed, thickness and material tokens are parsed, sequences containing prohibited materials are rejected, and duplicates are removed using the physical structure representation. Every retained candidate is then simulated for both polarizations with the same TMM implementation used for labels.")
add_paragraph(doc, "The separation of proposal and verification has three benefits. First, it prevents a high-probability sequence from being accepted when its independently simulated spectrum is poor. Second, it exploits the one-to-many nature of thin-film inverse design: stochastic samples can occupy different basins of the physical design landscape. Third, it makes search cost measurable through the number of requested candidates and TMM calls. The strict OOD experiment later quantifies this tradeoff for K=1,16,64.")
add_paragraph(doc, "3.5 Architecture semantics and reproducible loading", style="Heading 2")
add_paragraph(doc, "A practical issue discovered during migration was that parameter shape alone was insufficient to identify computation semantics. Historical OptoGPT checkpoints used fully connected and feed-forward paths that were shape-compatible with a later implementation but differed in activation behavior. Loading such weights into the wrong forward definition would not trigger a state-dictionary error, yet it would change the function represented by the checkpoint. The joint implementation therefore records an architecture version, and the formal model uses joint_sp_legacy_v1 semantics. Known legacy checkpoints are identified through metadata and complete hashes; unknown unversioned checkpoints are rejected unless their semantics are explicitly supplied.")
add_paragraph(doc, "Although checkpoint governance is not the central optical innovation, it supports reproducibility because the reported spectra depend on the exact computation graph, not merely the parameter tensor dimensions. The final checkpoint also records the pretrained-weight hash and dataset manifest, allowing the model, data contract, and training history to be associated unambiguously.")

add_paragraph(doc, "4. Joint Dataset and Training Protocol", style="Heading 1")
add_paragraph(doc, "4.1 Structure-paired data construction", style="Heading 2")
add_paragraph(doc, "The formal dataset was built from 500,000 legal dielectric structures. Each structure was simulated at 60° for both polarizations on the same 71-point wavelength grid. The two 142D channel records were concatenated according to Eq. (1). Data construction included finite-value checks, token parsing, material and thickness validation, layer-count validation, and sampled TMM recomputation. A completion marker was published only after these checks passed.")
add_paragraph(doc, "A SHA-256 physical-structure hash was used for deterministic splitting and duplicate control. The final counts were 400,006 training structures, 50,104 development structures, and 49,890 test structures. No structure-level leakage was detected across splits. Hashing at the structure level is more appropriate than comparing floating-point spectra because physically identical token sequences should remain in one split even if serialization or numerical formatting changes.")
add_table(doc, "Table 1", "Formal joint dataset and training configuration.", ["Category", "Setting", "Value"], [
    ["Dataset", "Training / development / test", "400,006 / 50,104 / 49,890"],
    ["Optical condition", "Spectrum order and dimension", "[Rs,Ts,Rp,Tp], 284D"],
    ["Spectral grid", "Wavelengths", "400–1100 nm, 10 nm interval, 71 points"],
    ["Geometry", "Training incidence", "60°"],
    ["Materials", "Allowed output vocabulary", "10 dielectrics"],
    ["Decoder", "Layers / model width / heads", "6 / 1024 / 8"],
    ["Optimization", "Formal epochs / batch size", "10 / 16"],
    ["Regularization", "Dropout / label smoothing", "0.1 / enabled"],
    ["Training system", "Precision / GPU", "mixed precision / RTX 4090 D"],
])
add_paragraph(doc, "4.2 Formal training protocol", style="Heading 2")
add_paragraph(doc, "The formal joint model was initialized from the approved OptoGPT checkpoint and trained for ten epochs. Adam optimization, label smoothing, mixed precision, differential learning rates, and ReduceLROnPlateau scheduling were used. The first two epochs served as the fusion warm-up, followed by full fine-tuning. Training loss decreased from 4.6647 at epoch 1 to 3.3164 at epoch 10, and development loss decreased from 4.3498 to 3.0381. The minimum development loss occurred at epoch 10.")
add_paragraph(doc, "The loss is token-level cross entropy, not optical reconstruction error. It evaluates the probability assigned to the ground-truth sequence under teacher forcing. Optical error is obtained only after free-running generation and TMM recomputation. Maintaining this distinction is necessary because a lower token loss does not imply that the model has recovered the original structure or that one decoded sequence is the best physical solution.")
add_paragraph(doc, "4.3 Development trajectory versus formal protocol", style="Heading 2")
add_paragraph(doc, "Figure 2(a) also includes v3, v4, and ultimate checkpoints that document the broader development trajectory. These stages are shown sequentially for interpretability but are not treated as one continuous 27-epoch formal run. The ten-epoch history above is the authoritative protocol for the formal 500k model. The v3/v4 channel diagnostics in Fig. 2(b) are retained to visualize the persistent s/p asymmetry during development, rather than to define the final test metric.")
add_paragraph(doc, "The material-vocabulary panel illustrates the transition from an 18-material source vocabulary containing eight metals or semiconductors to a ten-dielectric output contract. The flow diagram then summarizes the final evaluation path: joint 284D condition, shared-structure decoding, dielectric masking, exact 60° TMM, and optical and structural statistics.")
add_paragraph(doc, "4.4 Checkpointing and computational reproducibility", style="Heading 2")
add_paragraph(doc, "Each training checkpoint stores the model, optimizer, scheduler, AMP scaler, epoch, global step, Python/NumPy/PyTorch/CUDA random states, architecture version, pretrained hash, and data-manifest information. Resume behavior is defined at epoch boundaries. The original pretrained checkpoint remains unchanged. Windows deployment and console-specific fixes were required to complete the formal run, but these platform details are implementation support rather than algorithmic contributions and are best documented in supplementary reproducibility material.")
add_figure(doc, "figure2.png", "Figure 2", "Training progression and quantitative diagnostics of the joint 284D s+p model. (a) Archived joint-model development trajectory, with v3 and v4 shown as successive development stages and the ultimate checkpoint as the final recorded state. The formal 500k protocol is a separate ten-epoch run described in the text. (b) Channel-separated total-MAE diagnostics within the same shared-structure joint model, showing the larger s-channel error. (c) Purification from an 18-material source vocabulary to ten dielectric output materials. (d) Joint training and physics-evaluation contract. The material mask is a deterministic legality constraint; no same-checkpoint mask-off MAE ablation is claimed.")

add_paragraph(doc, "5. Joint-Model Results", style="Heading 1")
add_paragraph(doc, "5.1 Primary held-out evaluation", style="Heading 2")
add_paragraph(doc, "The primary quantitative evaluation sampled 100 targets from the 49,890-structure test pool. All 100 targets produced a valid decoded structure, and all 100 structures completed s/p TMM evaluation. The mean errors were Es=0.05084, Ep=0.02145, and Ejoint=0.03615. The difference between Es and Ep confirms that the s channel is the dominant source of joint error at 60°, consistent with its larger Fresnel reflection away from the Brewster condition.")
add_table(doc, "Table 2", "Formal held-out performance of JP-OptoGPT on 100 sampled test targets.", ["Metric", "Value", "Interpretation"], [
    ["Valid decoding", "100/100", "All sampled targets produced legal sequences"],
    ["TMM success", "100/100", "All decoded structures completed s/p evaluation"],
    ["Mean Es", "0.05084", "s-channel spectral MAE"],
    ["Mean Ep", "0.02145", "p-channel spectral MAE"],
    ["Mean Ejoint", "0.03615", "Equal-weight joint error"],
    ["Average mean Ts", "0.50367", "Mixed target distribution, not a high-T-only set"],
    ["Average mean Tp", "0.90179", "Mixed target distribution"],
    ["Average p05 Ts / Tp", "0.25134 / 0.67394", "Lower-tail transmission over wavelength"],
])
add_paragraph(doc, "The transmission averages in Table 2 should not be interpreted as the maximum antireflection performance because the randomly sampled test pool contains diverse optical targets. They primarily verify that the model reconstructs the requested joint spectra and reveal polarization asymmetry. The dedicated high-transmission application is evaluated separately in Sec. 7.")
add_paragraph(doc, "5.2 Auxiliary statistical evaluation", style="Heading 2")
add_paragraph(doc, "Figure 3 summarizes an auxiliary 200-sample statistical record associated with the final joint checkpoint family. All 200 structures are TMM-valid. The mean total MAE is 0.03157 with a standard deviation of 0.02157. The generated layer counts range from 1 to 15 with a mean of 6.50. The model uses all ten allowed dielectric materials, with TiO2, ZnO, Ta2O5, Si3N4, and MgF2 appearing most frequently in the retained structures.")
add_paragraph(doc, "This 200-sample record provides distributional and structural context, whereas the formal 100-sample result in Table 2 remains the primary quantitative claim because its sampling contract, joint condition, and completion marker are explicitly archived. We avoid combining the two sample counts into a single estimate. Agreement between their overall error scales nevertheless supports the stability of the reported reconstruction behavior.")
add_paragraph(doc, "5.3 Representative reconstructions", style="Heading 2")
add_paragraph(doc, "The best validation example in Fig. 3(a) is nearly indistinguishable from its target over the full band, illustrating that a simple generated structure can reproduce a smooth spectrum. The typical example in Fig. 3(b) contains sharper spectral features and exhibits visible local deviations, especially where the target reflectance changes rapidly. These cases show why a single aggregate MAE is insufficient: similar average error can arise from broad small deviations or from a few narrow spectral mismatches.")
add_paragraph(doc, "The wavelength-resolved error in Fig. 3(f) remains approximately 0.02–0.04 across much of the band and increases near portions of the long-wavelength range. Because reflection and transmission are coupled by energy conservation for the dielectric stack, their error profiles are strongly related. The statistics should be interpreted as optical reconstruction behavior of the generated structures, not as exact recovery of the data-generating token sequence.")
add_paragraph(doc, "5.4 Structural diversity and inverse non-uniqueness", style="Heading 2")
add_paragraph(doc, "The broad layer-count distribution and material usage support the generative interpretation of JP-OptoGPT. The model is not restricted to one fixed stack family and can propose structures with different complexities. This flexibility is valuable for fabrication-oriented selection because two candidates with similar spectra may differ in total thickness, number of interfaces, available materials, or sensitivity to deposition error.")
add_paragraph(doc, "Generative diversity is not automatically useful, however. A collection of different token sequences can contain invalid, redundant, or optically poor candidates. The logits mask, physical hashing, and TMM reranking convert raw sequence diversity into a set of legal, physically comparable designs. The strict OOD analysis in the next section quantifies how additional candidates translate into lower optical error.")
add_figure(doc, "figure3.png", "Figure 3", "Spectral reconstruction and structural statistics for the joint 284D model. (a,b) Best and typical TMM-recomputed validation examples. Gray curves denote the shared target and dashed coral/teal curves denote the generated s/p responses. (c) Total-MAE distribution over 200 valid auxiliary samples (mean 0.03157, standard deviation 0.02157). (d) Material-token frequency. (e) Generated layer-count distribution (range 1–15, mean 6.50). (f) Wavelength-resolved channel error. The formal 100-sample evaluation in Table 2 is used as the primary quantitative result.")

add_paragraph(doc, "6. Candidate-Budget and OOD Analysis", style="Heading 1")
add_paragraph(doc, "6.1 Strict OOD construction", style="Heading 2")
add_paragraph(doc, "An inverse model can appear accurate when the evaluation reproduces the structural grammar used for training. We therefore constructed a stricter 60-target OOD set. The structures contain 15–20 layers, thicknesses distributed continuously between 11 and 499 nm rather than restricted to the 10 nm grid, and four structural families: random, graded, strongly alternating, and bimodal-thickness stacks. Twelve targets additionally include small spectral perturbations. No exact structure overlap was found against the locally verifiable training structures.")
add_paragraph(doc, "The OOD targets remain physically meaningful joint spectra because both polarization channels are produced from the same structure. The test changes the structural distribution rather than replacing the forward physics. This makes it possible to evaluate whether the decoder and reranking process can propose spectrally equivalent structures outside the discrete generation grammar.")
add_paragraph(doc, "6.2 Candidate-budget dependence", style="Heading 2")
add_paragraph(doc, "Table 3 compares one greedy candidate with 16 and 64 sampled candidates followed by exact TMM ranking. With K=1, the mean joint MAE is 0.0930. Increasing the budget to 16 reduces the mean to 0.0719, and K=64 reduces it to 0.0634. The median and worst-case errors decrease as well. Relative to one candidate, 64-candidate reranking reduces the mean error by approximately 31.8%.")
add_table(doc, "Table 3", "Strict OOD performance as a function of candidate budget.", ["Inference strategy", "Mean Ejoint", "Median Ejoint", "Worst Ejoint"], [
    ["1 candidate, greedy", "0.0930", "0.0889", "0.1722"],
    ["16 candidates + TMM reranking", "0.0719", "0.0714", "0.1300"],
    ["64 candidates + TMM reranking", "0.0634", "0.0643", "0.1131"],
])
add_paragraph(doc, "The comparison provides a formal baseline for the physical reranking component. The K=1 result represents a joint model without a multi-candidate selection advantage; K=16 and K=64 quantify the benefit of proposing additional structures and selecting them using the same physical objective. It would be inaccurate to describe the 64-candidate number as the error of one direct model output. The result instead characterizes the full JP-OptoGPT inference framework at a specified search budget.")
add_paragraph(doc, "The improvement exhibits diminishing returns: the first increase from 1 to 16 candidates yields a larger gain than the increase from 16 to 64. This is expected when the decoder probability mass concentrates around a limited number of structural modes. Future work should report error as a function of cumulative TMM calls and investigate whether diversity-aware sampling or active acquisition can reach the same accuracy with fewer evaluations.")
add_paragraph(doc, "6.3 Polarization and structural-family effects", style="Heading 2")
add_paragraph(doc, "At K=64, the mean s-channel MAE is 0.0827, whereas the p-channel mean is 0.0440. The same asymmetry observed in the held-out set therefore persists under distribution shift. Strongly alternating structures are the hardest family, with a mean joint MAE of approximately 0.0826 at 64 candidates. Such sequences produce rapid impedance and phase changes and are less likely to be represented by a short, high-probability output sequence.")
add_paragraph(doc, "The exact original structure is recovered for none of the 60 OOD targets. This is not evidence that the optical inverse problem has failed. Instead, it confirms that the model finds different structures with approximately equivalent spectra, which is consistent with the non-uniqueness of multilayer synthesis and with prior probabilistic inverse-design studies [7,11]. Structural recovery and spectral inverse design are different objectives; the present work optimizes the latter.")
add_paragraph(doc, "6.4 Wide-angle diagnostic", style="Heading 2")
add_paragraph(doc, "Although JP-OptoGPT is trained at 60°, generated high-transmission candidates were scanned from 0° to 80° using exact TMM. Nineteen unique candidates were collected, and 16 completed dense 1° evaluation. No candidate satisfied the strict mean, lower-tail, and minimum-transmission gates over the full angular range. A representative structure provides approximately equal mean transmission of 0.937 near normal incidence but degrades to mean Ts≈0.3525 and mean Tp≈0.5925 at 80°.")
add_paragraph(doc, "This diagnostic defines the domain of the current model. A 60° condition can produce a useful design at or near its training angle, but it should not be described as a 0–75° omnidirectional coating. Wide-angle broadband antireflection is known to require careful control of low-index matching, material dispersion, and thickness tolerance [12]. Multi-angle conditioning, worst-angle objectives, or an expanded refractive-index space are required before a wide-angle claim can be supported.")

add_paragraph(doc, "7. Finite-Glass High-Transmission Application", style="Heading 1")
add_paragraph(doc, "7.1 Model-generated finite-glass candidates", style="Heading 2")
add_paragraph(doc, "The application target is a flat joint response with Rs=Rp=0.05 and Ts=Tp=0.95 over 400–1100 nm. From 4096 requested sequences, 2032 unique TMM-valid candidates were retained. The three candidates in the upper part of Fig. 4 were selected using the limiting-polarization transmission, the mean unpolarized transmission, and structural diversity. Their material sequences and layer counts differ, demonstrating that the joint target admits multiple legal proposals.")
add_paragraph(doc, "Candidate A provides the best limiting-channel performance in the displayed set, with mean Ts=0.7733 and mean Tp=0.9818 on a 500 μm finite glass substrate. Candidate B is much thinner and achieves mean Tp≈0.993 but lower mean Ts≈0.748. Candidate C provides a different material sequence and mean Ts≈0.771 with mean Tp≈0.963. The p channel remains close to the high-transmission reference, while the s reflectance remains substantially above 5%.")
add_paragraph(doc, "The selected model candidate is nevertheless a successful joint application when compared with the physically relevant bare finite-glass baseline. Bare glass has mean Ts=0.6938 and mean Tp=0.9967, giving mean unpolarized transmission 0.8453. Candidate A raises mean Ts to 0.7733 and mean unpolarized transmission to 0.8776, while mean Tp decreases slightly to 0.9818. The result improves the limiting polarization and the balanced joint response rather than improving both channels independently.")
add_paragraph(doc, "7.2 Physics-refined double-sided extension", style="Heading 2")
add_paragraph(doc, "The lower part of Fig. 4 extends the generative workflow to a double-sided finite-glass design. This final structure is not claimed to be a one-step JP-OptoGPT output. Model-generated candidates provide physically informed initialization, after which the front and back coatings are evaluated in the finite-glass geometry and independently refined. The coherent coatings and incoherent substrate are recomputed together throughout the refinement.")
add_paragraph(doc, "The optimized front coating is MgF2 143.9 nm / Al2O3 151.5 nm / TiO2 10.0 nm / SiO2 32.9 nm, and the back coating is MgO 169.3 nm / Al2O3 108.9 nm / MgF2 139.1 nm. It reaches mean Rs=0.1634, mean Rp=0.0197, mean Ts=0.8350, and mean Tp=0.9792. Mean unpolarized transmission is 0.9071. Relative to bare glass, mean Ts increases by 0.1412 absolute and mean unpolarized transmission by 0.0618. Relative to the best displayed model candidate, refinement provides a further 0.0617 increase in mean Ts and a 0.0295 increase in mean unpolarized transmission.")
add_table(doc, "Table 4", "Finite-glass application progression at 60° over 400–1100 nm.", ["Design", "Mean Ts", "Mean Tp", "Mean Tunpol", "Interpretation"], [
    ["Bare finite glass", "0.6938", "0.9967", "0.8453", "Physical application baseline"],
    ["Best model-generated candidate", "0.7733", "0.9818", "0.8776", "Improved limiting s channel"],
    ["Physics-refined double-sided design", "0.8350", "0.9792", "0.9071", "Further joint-performance improvement"],
])
add_paragraph(doc, "7.3 Success criterion and remaining limitation", style="Heading 2")
add_paragraph(doc, "The application success is defined by three verified outcomes: JP-OptoGPT produces legal high-transmission dielectric candidates; multi-candidate physical selection identifies a structure that improves the limiting s channel and mean unpolarized transmission relative to bare finite glass; and the generative workflow supplies useful initialization for a stronger double-sided design. The result is not defined as full attainment of the ideal target.")
add_paragraph(doc, "Neither the selected model candidate nor the double-sided refined design simultaneously reaches mean Ts≥0.95 and mean Tp≥0.95. The double-sided result also fails the stricter wavelength-wise reflection and robustness gates, and the available perturbation study does not establish fabrication readiness. The example is therefore a computational, TMM-validated engineering demonstration. It supports the utility of the framework while exposing the persistent s-polarized limitation.")
add_paragraph(doc, "The progression in Table 4 is physically informative. Bare glass is already nearly transparent for p polarization near the Brewster regime, leaving little room for improvement. The design problem is dominated by the s channel and by the rear-surface penalty. A successful joint objective must therefore accept a small decrease in p transmission if that trade yields a larger increase in s transmission and average throughput.")
add_figure(doc, "figure4_extended.png", "Figure 4", "Finite-glass high-transmission application and physics-refined double-sided extension. The original three-candidate Figure 4 is preserved separately. (a–c) Model-generated candidates selected from 2032 unique TMM-valid structures obtained from 4096 requests under the same flat joint target. Their distinct stacks demonstrate multiple legal proposals. (d) Physics-refined engineering extension initialized from the joint generative workflow. The lower comparison shows bare finite glass, the best displayed model candidate, and the refined double-sided design. The limiting s-polarized mean transmission increases from 0.6938 to 0.7733 and then to 0.8350; mean unpolarized transmission increases from 0.8453 to 0.8776 and 0.9071. Mean Tp decreases slightly from 0.9967 to 0.9818 and 0.9792. The ideal dual-polarization 0.95 target remains unmet.")

add_paragraph(doc, "8. Discussion", style="Heading 1")
add_paragraph(doc, "8.1 What is added beyond OptoGPT", style="Heading 2")
add_paragraph(doc, "The original OptoGPT establishes a general sequence-generation paradigm for multilayer structures [8]. JP-OptoGPT adds a specialized physical contract and inference system. The first addition is the structure-paired joint condition, which changes the task from single-channel generation to one shared design constrained by four spectra. The second is polarization-aware encoding and fusion. The third is deterministic material legality. The fourth is candidate-level exact TMM reranking, evaluated explicitly as a function of search budget. The fifth is transfer from a semi-infinite generation setting to finite-glass and double-sided engineering analysis.")
add_paragraph(doc, "These additions are mutually dependent. A joint condition without exact validation could still generate a spectrum mismatch. TMM reranking without material masking could select absorbing or prohibited candidates. A dielectric mask without a structure-paired dataset would not solve the shared-polarization task. The main contribution is therefore the integrated physics-validated workflow rather than any one module in isolation.")
add_paragraph(doc, "8.2 Comparison with alternative learning paradigms", style="Heading 2")
add_paragraph(doc, "Reinforcement-learning approaches can build multilayers sequentially and optimize a physical reward directly [5,6]. Their advantage is close coupling between actions and optical objectives, but a new search or policy adaptation may be required for different target families. Mixture-density and invertible models explicitly represent conditional multiplicity [7,11], although many demonstrations use fixed structural dimensionality. JP-OptoGPT instead inherits a variable-length token grammar and uses sampling plus TMM to expose multiple solutions.")
add_paragraph(doc, "Benchmarking studies emphasize that no inverse architecture is uniformly best across all photonic problems [13]. The present results support a hybrid perspective: the Transformer supplies a learned structural prior, while the forward solver determines physical quality. This division is especially appropriate for multilayer films because TMM is fast, exact within its assumptions, and differentiates subtle s/p effects that token likelihood cannot encode as a hard guarantee.")
add_paragraph(doc, "8.3 Data efficiency and active-learning opportunity", style="Heading 2")
add_paragraph(doc, "The formal dataset contains 500,000 structures, yet the hardest high-angle, high-transmission region remains sparsely represented. Uniform random sampling spends most TMM calls on ordinary structures rather than on the narrow region where both polarization channels perform well. Deep active learning has been shown to improve data-driven thin-film design by selecting informative simulations instead of blindly expanding the dataset [14], and uncertainty-guided sampling is widely used in materials discovery [15].")
add_paragraph(doc, "A future JP-OptoGPT loop could combine forward-model uncertainty, target performance, and structural diversity. Candidate structures with high predicted transmission, high ensemble disagreement, or novel material sequences would be labeled by TMM and added to the training set. Comparisons should use cumulative TMM calls as the horizontal axis and include random sampling, performance-only selection, and the current self-improving workflow as baselines. The present work includes the record and deduplication infrastructure for such a study but does not claim a completed active-learning gain.")
add_paragraph(doc, "8.4 Physical and fabrication limitations", style="Heading 2")
add_paragraph(doc, "TMM assumes planar, laterally homogeneous, isotropic layers with known optical constants. Surface roughness, interdiffusion, deposition-induced density changes, and anisotropy are not modeled. The material database itself is a source of uncertainty because refractive indices depend on wavelength, process, and film microstructure. A design that is optimal for one tabulated dataset may shift after fabrication.")
add_paragraph(doc, "The formal token grammar discretizes thickness, whereas deposition errors are continuous. The double-sided refinement partially addresses this by using continuous thickness optimization, but the reported perturbation results do not satisfy the strict manufacturing gate. Future work should include distributionally robust objectives, correlated thickness error, material-index uncertainty, and experimentally calibrated optical constants. Minimum layer thickness and total coating thickness should also be included directly in the ranking objective.")
add_paragraph(doc, "The angle scan shows that the current model is not wide-angle. Achieving broadband dual-polarization transmission at 70–80° may require porous or graded low-index layers that are unavailable in the ten-material dense vocabulary. Wide-angle antireflection coatings fabricated by atomic layer deposition illustrate the value of controlled index matching and material engineering [12]. Expanding the token vocabulary to effective-medium or continuously parameterized materials is a promising route, but it changes both the fabrication assumptions and the learning space.")
add_paragraph(doc, "8.5 Evaluation boundaries", style="Heading 2")
add_paragraph(doc, "Several reporting boundaries are maintained throughout the paper. First, optical MAE is not called structure accuracy. Second, a best-of-K result is always identified as multi-candidate performance. Third, the auxiliary 200-sample statistics are not merged with the formal 100-sample test result. Fourth, the material mask is described as a legality guarantee rather than an unsupported accuracy ablation. Fifth, the physics-refined double-sided design is distinguished from a direct model output. Finally, all results are computational; no measured or fabricated device is claimed.")
add_paragraph(doc, "These distinctions are important for assessing generative inverse design. A method can be useful even when it does not recover the hidden structure or attain an ideal target, provided it generates verifiable candidates that improve a relevant physical baseline. Here, the formal test, OOD candidate-budget study, finite-glass comparison, and wide-angle failure analysis together give a more complete picture than a single favorable spectrum.")

add_paragraph(doc, "9. Conclusion", style="Heading 1")
add_paragraph(doc, "We developed JP-OptoGPT, a physics-validated extension of OptoGPT for shared-structure joint s- and p-polarization inverse design. The method combines a 284D structure-paired condition, polarization-aware encoding, a transferred variable-length autoregressive decoder, deterministic dielectric masking, and exact multi-candidate TMM reranking. A 500,000-structure dataset with deterministic hashing and zero detected split leakage supports formal training and evaluation.")
add_paragraph(doc, "On 100 held-out targets, JP-OptoGPT achieves 100% valid decoding and TMM success with mean Ejoint=0.03615. Strict OOD evaluation shows that increasing the candidate budget from one to 64 reduces mean joint MAE from 0.0930 to 0.0634, quantifying the value of physical reranking. In the finite-glass application, the best displayed model candidate raises mean Ts from 0.6938 for bare glass to 0.7733 and mean unpolarized transmission from 0.8453 to 0.8776. Physics-refined double-sided optimization further raises these values to 0.8350 and 0.9071.")
add_paragraph(doc, "The results also identify the present boundary: s polarization remains the limiting channel, the ideal dual-polarization 0.95 transmission target is not fully attained, fabrication robustness is not established, and a 60° model does not generalize to the full 0–80° range. By reporting both the improvements and these limitations, JP-OptoGPT provides a reproducible baseline for future multi-angle, materials-expanded, robust, and active-learning thin-film design.")

add_paragraph(doc, "References", style="Heading 1")
references = [
    "H. A. Macleod, Thin-Film Optical Filters, 5th ed. (CRC Press, 2017).",
    "S. So, T. Badloe, J. Noh, J. Rho, and J. Bravo-Abad, “Deep learning enabled inverse design in nanophotonics,” Nanophotonics 9, 1041–1057 (2020), doi:10.1515/nanoph-2019-0474.",
    "Z. Liu, D. Zhu, S. P. Rodrigues, K.-T. Lee, and W. Cai, “Generative model for the inverse design of metasurfaces,” Nano Lett. 18, 6570–6576 (2018), doi:10.1021/acs.nanolett.8b03171.",
    "P. R. Wiecha, A. Arbouet, C. Girard, and O. L. Muskens, “Deep learning in nano-photonics: inverse design and beyond,” Photon. Res. 9, B182–B200 (2021), doi:10.1364/PRJ.415960.",
    "H. Wang, Z. Zheng, C. Ji, and L. J. Guo, “Automated multi-layer optical design via deep reinforcement learning,” Mach. Learn.: Sci. Technol. 2, 025013 (2021), doi:10.1088/2632-2153/abc327.",
    "A.-Q. Jiang, O. Yoshie, and L.-Y. Chen, “Multilayer optical thin film design with deep Q learning,” Sci. Rep. 10, 12780 (2020), doi:10.1038/s41598-020-69754-w.",
    "R. Unni, K. Yao, and Y. Zheng, “Deep convolutional mixture density network for inverse design of layered photonic structures,” ACS Photonics 7, 2703–2712 (2020), doi:10.1021/acsphotonics.0c00630.",
    "T. Ma, H. Wang, and L. J. Guo, “OptoGPT: a foundation model for inverse design in optical multilayer thin film structures,” Opto-Electron. Adv. 7, 240062 (2024), doi:10.29026/oea.2024.240062.",
    "A. Luce, A. Mahdavi, F. Marquardt, and H. Wankerl, “TMM-Fast, a transfer matrix computation package for multilayer thin-film optimization: tutorial,” J. Opt. Soc. Am. A 39, 1007–1013 (2022), doi:10.1364/JOSAA.450928.",
    "S. J. Byrnes, “Multilayer optical calculations,” arXiv:1603.02720 (2016), doi:10.48550/arXiv.1603.02720.",
    "A. Luce, A. Mahdavi, H. Wankerl, and F. Marquardt, “Investigation of inverse design of multilayer thin-films with conditional invertible neural networks,” Mach. Learn.: Sci. Technol. 4, 015014 (2023), doi:10.1088/2632-2153/acb48d.",
    "K. Pfeiffer, L. Ghazaryan, U. Schulz, and A. Szeghalmi, “Wide-angle broadband antireflection coatings prepared by atomic layer deposition,” ACS Appl. Mater. Interfaces 11, 21887–21894 (2019), doi:10.1021/acsami.9b03125.",
    "T. Ma, M. Tobah, H. Wang, and L. J. Guo, “Benchmarking deep learning-based models on nanophotonic inverse design problems,” Opto-Electron. Sci. 1, 210012 (2022), doi:10.29026/oes.2022.210012.",
    "Y. Hong and D. P. Nicholls, “Data-driven design of thin-film optical systems using deep active learning,” Opt. Express 30, 22901–22916 (2022), doi:10.1364/OE.459295.",
    "T. Lookman, P. V. Balachandran, D. Xue, and R. Yuan, “Active learning in materials science with emphasis on adaptive sampling using uncertainties for targeted design,” npj Comput. Mater. 5, 21 (2019), doi:10.1038/s41524-019-0153-8.",
]
for index, item in enumerate(references, 1):
    add_reference(doc, index, item)

doc.save(OUT)

text = []
for paragraph in doc.paragraphs:
    text.append(paragraph.text)
for table in doc.tables:
    for row in table.rows:
        text.extend(cell.text for cell in row.cells)
word_count = len(re.findall(r"\b[\w’'-]+\b", "\n".join(text)))
print(OUT)
print("WORD_COUNT", word_count)
