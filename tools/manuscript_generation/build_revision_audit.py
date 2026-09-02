from collections import Counter
from datetime import date
from hashlib import sha256
import csv
import json
from pathlib import Path
import re
from xml.etree import ElementTree as ET
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parent
WORK = ROOT / "paper_rewriting_output"
AUDIT = WORK / "scientific_audit"
ORIGINAL = ROOT / "Physics_Validated_JP_OptoGPT_Manuscript_Body.docx"
REVISED = ROOT / "Physics_Validated_JP_OptoGPT_Manuscript_Body_Revised.docx"


PUBLIC_SOURCES = [
    ("E001", "Thin-Film Optical Filters, 5th ed.", ["H. A. Macleod"], 2017, "", "book"),
    ("E002", "Deep learning enabled inverse design in nanophotonics", ["Sunae So", "Trevon Badloe", "Jaebum Noh", "Junsuk Rho", "Jorge Bravo-Abad"], 2020, "10.1515/nanoph-2019-0474", "journal_article"),
    ("E003", "Generative Model for the Inverse Design of Metasurfaces", ["Zhaocheng Liu", "Dayu Zhu", "Sean P. Rodrigues", "Kyu-Tae Lee", "Wenshan Cai"], 2018, "10.1021/acs.nanolett.8b03171", "journal_article"),
    ("E004", "Deep learning in nano-photonics: inverse design and beyond", ["Peter R. Wiecha", "Arnaud Arbouet", "Christian Girard", "Otto L. Muskens"], 2021, "10.1364/PRJ.415960", "journal_article"),
    ("E005", "Automated multi-layer optical design via deep reinforcement learning", ["Haozhu Wang", "Zeyu Zheng", "Chengang Ji", "L. Jay Guo"], 2021, "10.1088/2632-2153/abc327", "journal_article"),
    ("E006", "Multilayer optical thin film design with deep Q learning", ["An-Qing Jiang", "Osamu Yoshie", "Liang-Yao Chen"], 2020, "10.1038/s41598-020-69754-w", "journal_article"),
    ("E007", "Deep Convolutional Mixture Density Network for Inverse Design of Layered Photonic Structures", ["Rohit Unni", "Kan Yao", "Yuebing Zheng"], 2020, "10.1021/acsphotonics.0c00630", "journal_article"),
    ("E008", "OptoGPT: A foundation model for inverse design in optical multilayer thin film structures", ["Taigao Ma", "Haozhu Wang", "L. Jay Guo"], 2024, "10.29026/oea.2024.240062", "journal_article"),
    ("E009", "TMM-Fast, a transfer matrix computation package for multilayer thin-film optimization: tutorial", ["Alexander Luce", "Ali Mahdavi", "Florian Marquardt", "Heribert Wankerl"], 2022, "10.1364/JOSAA.450928", "journal_article"),
    ("E010", "Multilayer optical calculations", ["Steven J. Byrnes"], 2016, "10.48550/arXiv.1603.02720", "preprint"),
    ("E011", "Investigation of inverse design of multilayer thin-films with conditional invertible neural networks", ["Alexander Luce", "Ali Mahdavi", "Heribert Wankerl", "Florian Marquardt"], 2023, "10.1088/2632-2153/acb48d", "journal_article"),
    ("E012", "Wide-Angle Broadband Antireflection Coatings Prepared by Atomic Layer Deposition", ["Kristin Pfeiffer", "Lilit Ghazaryan", "Ulrike Schulz", "Adriana Szeghalmi"], 2019, "10.1021/acsami.9b03125", "journal_article"),
    ("E013", "Benchmarking deep learning-based models on nanophotonic inverse design problems", ["Taigao Ma", "Mustafa Tobah", "Haozhu Wang", "L. Jay Guo"], 2022, "10.29026/oes.2022.210012", "journal_article"),
    ("E014", "Data-driven design of thin-film optical systems using deep active learning", ["Youngjoon Hong", "David P. Nicholls"], 2022, "10.1364/OE.459295", "journal_article"),
    ("E015", "Active learning in materials science with emphasis on adaptive sampling using uncertainties for targeted design", ["Turab Lookman", "Prasanna V. Balachandran", "Dezhen Xue", "Ruihao Yuan"], 2019, "10.1038/s41524-019-0153-8", "journal_article"),
]

INTERNAL_SOURCES = [
    ("E016", "Formal 500k structure manifest and split audit", "Sec. 4; Table 1"),
    ("E017", "Formal ten-epoch training history", "Sec. 4.2; Fig. 2"),
    ("E018", "Formal 100-target held-out evaluation", "Sec. 5.1; Table 2"),
    ("E019", "Auxiliary 200-sample reconstruction archive", "Sec. 5.2; Fig. 3"),
    ("E020", "Strict 60-target OOD candidate-budget archive", "Sec. 6; Table 3"),
    ("E021", "Dense 0-80 degree angular scan archive", "Sec. 6.4"),
    ("E022", "Finite-glass candidate archive", "Sec. 7.1; Fig. 4"),
    ("E023", "Double-sided finite-glass refinement archive", "Sec. 7.2; Table 4; Fig. 4"),
]

CLAIMS = [
    ("C001", "Introduction", "factual", "OptoGPT represents multilayer structures as conditional autoregressive material-thickness token sequences.", "E008"),
    ("C002", "Methods", "method", "The formal dataset contains 500000 structure-paired joint-polarization records split by physical hash.", "E016"),
    ("C003", "Methods", "method", "The formal training protocol contains ten epochs with a two-epoch fusion warm-up.", "E017"),
    ("C004", "Results", "result", "The held-out evaluation completed 100 legal decodes and 100 TMM evaluations with mean joint error 0.03615.", "E018"),
    ("C005", "Results", "result", "The auxiliary 200-sample record has mean total MAE 0.03157 and remains separate from the formal test.", "E019"),
    ("C006", "Results", "result", "Increasing the OOD candidate budget from one to 64 reduces mean joint MAE from 0.0930 to 0.0634.", "E020"),
    ("C007", "Results", "result", "No candidate passes the full 0-80 degree transmission gate.", "E021"),
    ("C008", "Application", "result", "The selected model candidate raises mean s transmission from 0.6938 for bare glass to 0.7733.", "E022"),
    ("C009", "Application", "result", "Double-sided physical refinement raises mean s transmission to 0.8350 and unpolarized mean transmission to 0.9071.", "E023"),
    ("C010", "Discussion", "interpretive", "The refined double-sided coating is not a one-step JP-OptoGPT output.", "E023"),
    ("C011", "Discussion", "factual", "Active learning has been applied to thin-film optical-system design.", "E014"),
    ("C012", "Discussion", "factual", "Uncertainty-guided adaptive sampling is established in materials design.", "E015"),
]


def source_manifest():
    sources = []
    for evidence_id, title, authors, year, doi, source_type in PUBLIC_SOURCES:
        sources.append({
            "authors": authors,
            "confidentiality": "public",
            "evidence_id": evidence_id,
            "identifiers": {"doi": doi, "isbn": "", "pmcid": "", "pmid": "", "url": ""},
            "locator": "Bibliographic metadata checked on 2026-08-15; full-text proposition locator requires author verification.",
            "source_type": source_type,
            "title": title,
            "verification": {"source_opened": False, "status": "unverified", "verified_by": "", "verified_on": ""},
            "year": year,
        })
    for evidence_id, title, locator in INTERNAL_SOURCES:
        sources.append({
            "authors": [],
            "confidentiality": "restricted",
            "evidence_id": evidence_id,
            "identifiers": {"doi": "", "isbn": "", "pmcid": "", "pmid": "", "url": ""},
            "locator": locator,
            "source_type": "other",
            "title": title,
            "verification": {"source_opened": False, "status": "unverified", "verified_by": "", "verified_on": ""},
            "year": 2026,
        })
    (AUDIT / "source_manifest.json").write_text(json.dumps({"schema_version": "1.0", "sources": sources}, indent=2), encoding="utf-8")


def claims_csv():
    with (AUDIT / "claims.csv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(["claim_id", "section", "claim_kind", "claim_text_sha256", "evidence_ids", "verification_status", "uncertainty", "analysis_intent"])
        for claim_id, section, kind, claim, evidence in CLAIMS:
            digest = sha256(" ".join(claim.lower().split()).encode()).hexdigest()
            writer.writerow([claim_id, section, kind, digest, evidence, "unverified", "bounded_by_manuscript", "descriptive"])


def consistency_manifest():
    methods = [
        {"analysis_intent": "descriptive", "method_id": "M001", "name": "Structure-paired s/p transfer-matrix evaluation", "outcome_ids": ["O001"], "protocol_status": "prespecified"},
        {"analysis_intent": "descriptive", "method_id": "M002", "name": "Multi-candidate TMM reranking", "outcome_ids": ["O002"], "protocol_status": "prespecified"},
        {"analysis_intent": "exploratory", "method_id": "M003", "name": "Finite-glass double-sided physical refinement", "outcome_ids": ["O003"], "protocol_status": "post_hoc"},
    ]
    facts = [
        ("N001", "formal dataset size", 500000, "structures", 500000, "E016", "Sec. 4"),
        ("N002", "held-out mean joint error", 0.03615, "dimensionless", 100, "E018", "Sec. 5.1"),
        ("N003", "OOD K1 mean joint MAE", 0.0930, "dimensionless", 60, "E020", "Sec. 6.2"),
        ("N004", "OOD K64 mean joint MAE", 0.0634, "dimensionless", 60, "E020", "Sec. 6.2"),
        ("N005", "bare-glass mean s transmission", 0.6938, "dimensionless", None, "E022", "Sec. 7.1"),
        ("N006", "model-candidate mean s transmission", 0.7733, "dimensionless", None, "E022", "Sec. 7.1"),
        ("N007", "refined mean s transmission", 0.8350, "dimensionless", None, "E023", "Sec. 7.2"),
        ("N008", "refined mean unpolarized transmission", 0.9071, "dimensionless", None, "E023", "Sec. 7.2"),
    ]
    numeric = [{
        "analysis_set": "reported project archive",
        "concept": concept,
        "denominator": None,
        "evidence_ids": [evidence],
        "fact_id": fact_id,
        "numerator": None,
        "sample_size": n,
        "section": section,
        "unit": unit,
        "value": value,
    } for fact_id, concept, value, unit, n, evidence, section in facts]
    results = [
        {"analysis_intent": "descriptive", "evidence_ids": ["E018"], "method_id": "M001", "outcome_id": "O001", "reported_sections": ["5.1", "9"], "result_id": "R001", "sample_size": 100},
        {"analysis_intent": "descriptive", "evidence_ids": ["E020"], "method_id": "M002", "outcome_id": "O002", "reported_sections": ["6.2", "9"], "result_id": "R002", "sample_size": 60},
        {"analysis_intent": "exploratory", "evidence_ids": ["E023"], "method_id": "M003", "outcome_id": "O003", "reported_sections": ["7.2", "9"], "result_id": "R003", "sample_size": 1},
    ]
    payload = {"methods": methods, "numeric_facts": numeric, "results": results, "schema_version": "1.0"}
    (AUDIT / "consistency_manifest.json").write_text(json.dumps(payload, indent=2), encoding="utf-8")


def document_text(path):
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs[:146])


def docx_structure(path):
    doc = Document(path)
    tables = [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in doc.tables
    ]
    references = [paragraph.text for paragraph in doc.paragraphs[147:162]]
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    }
    with ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
        media = {
            name: sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        }
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": tables,
        "references": references,
        "equations": len(root.findall(".//m:oMath", namespaces)),
        "drawings": len(root.findall(".//w:drawing", namespaces)),
        "image_refs": len(root.findall(".//a:blip", namespaces)),
        "media": media,
    }


def integrity_report():
    old = document_text(ORIGINAL)
    new = document_text(REVISED)
    old_structure = docx_structure(ORIGINAL)
    new_structure = docx_structure(REVISED)
    old_decimals = set(re.findall(r"\b\d+\.\d+\b", old))
    new_decimals = set(re.findall(r"\b\d+\.\d+\b", new))
    old_cites = Counter(re.findall(r"\[[0-9,\-\u2013]+\]", old))
    new_cites = Counter(re.findall(r"\[[0-9,\-\u2013]+\]", new))
    invariants = ["0.03615", "0.0930", "0.0634", "0.6938", "0.7733", "0.8350", "0.9071", "4.6647", "3.3164", "4.3498", "3.0381"]
    lines = [
        "# Revision Integrity Report", "",
        f"- Date: {date.today().isoformat()}",
        f"- Decimal values removed: {sorted(old_decimals - new_decimals)}",
        f"- Decimal values added: {sorted(new_decimals - old_decimals)}",
        f"- Citation marker multiset preserved: {'PASS' if old_cites == new_cites else 'FAIL'}",
        f"- Required numeric invariants present: {'PASS' if all(x in new for x in invariants) else 'FAIL'}",
        f"- Paragraph count preserved ({old_structure['paragraphs']}): {'PASS' if old_structure['paragraphs'] == new_structure['paragraphs'] else 'FAIL'}",
        f"- Table count and cell text preserved ({len(old_structure['tables'])}): {'PASS' if old_structure['tables'] == new_structure['tables'] else 'FAIL'}",
        f"- OMML equation count preserved ({old_structure['equations']}): {'PASS' if old_structure['equations'] == new_structure['equations'] else 'FAIL'}",
        f"- Drawing and image-reference counts preserved ({old_structure['drawings']}/{old_structure['image_refs']}): {'PASS' if (old_structure['drawings'], old_structure['image_refs']) == (new_structure['drawings'], new_structure['image_refs']) else 'FAIL'}",
        f"- Embedded image bytes preserved ({len(old_structure['media'])}): {'PASS' if old_structure['media'] == new_structure['media'] else 'FAIL'}",
        f"- Reference text and order preserved ({len(old_structure['references'])}): {'PASS' if old_structure['references'] == new_structure['references'] else 'FAIL'}",
        "- Full-text source proposition verification: PENDING AUTHOR REVIEW.",
        "",
    ]
    (WORK / "qc" / "revision_integrity_report.md").write_text("\n".join(lines), encoding="utf-8")


def main():
    source_manifest()
    claims_csv()
    consistency_manifest()
    integrity_report()
    print("AUDIT_MANIFESTS_WRITTEN", AUDIT)


if __name__ == "__main__":
    main()
