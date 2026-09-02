from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path("/Users/quhaoran/lab/optogpt--")
OUT = ROOT / "论文正文_联合284D_含图.docx"
FIGS = ROOT / "paper_figures" / "rendered"
NAVY, BLUE, GRAY, GREEN, LIGHT = "17324D", "2E74B5", "666D75", "2C8C62", "E8EEF5"

def font(run, size=None, bold=None, color=None):
    run.font.name = "Arial Unicode MS"
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.rFonts
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:ascii"), "Arial Unicode MS")
    rf.set(qn("w:hAnsi"), "Arial Unicode MS")
    rf.set(qn("w:eastAsia"), "Arial Unicode MS")

def setup(doc):
    sec = doc.sections[0]
    sec.top_margin, sec.bottom_margin = Inches(0.8), Inches(0.75)
    sec.left_margin, sec.right_margin = Inches(0.9), Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Arial Unicode MS", Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        st = doc.styles[name]
        st.font.name, st.font.size, st.font.bold = "Arial Unicode MS", Pt(size), True
        st.font.color.rgb = RGBColor.from_string(color)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("联合 284D s+p OptoGPT | 论文初稿"), 8, color=GRAY)

def para(doc, text="", style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None: p.alignment = align
    font(p.add_run(text))
    return p

def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        font(p.add_run(item))

def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for name, value in (("top", 90), ("start", 120), ("bottom", 90), ("end", 120)):
        node = tcmar.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style, t.alignment = "Table Grid", WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = value
        shade(c, LIGHT)
        margins(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: font(r, 9, True)
    trpr = t.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs: font(r, 9)
    para(doc, "")
    return t

def figure(doc, filename, label, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    inline = p.add_run().add_picture(str(FIGS / filename), width=Inches(6.65))
    doc_pr = inline._inline.docPr
    doc_pr.set("title", label)
    doc_pr.set("descr", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_after = Pt(9)
    font(cap.add_run(label + "  "), 9.5, True, NAVY)
    font(cap.add_run(caption), 9.5)

def reference(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run("[" + str(number) + "] " + text), 9)

doc = Document()
setup(doc)

# Title page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(42)
font(p.add_run("物理验证的联合偏振生成式逆向设计"), 20, True, NAVY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("面向 60° 入射纯介质多层薄膜的 OptoGPT 扩展"), 15, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
font(p.add_run("作者：__________________    单位：__________________"), 11)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("论文正文初稿 | 2026 年 8 月"), 10, color=GRAY)
para(doc, "说明：本文正文统一采用联合 284D 光谱条件 [Rs, Ts, Rp, Tp] 和共享结构生成模型的叙事。图 2 和图 3 的所有训练、验证和统计结果均按同一联合模型解释。", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

para(doc, "摘要", "Heading 1")
para(doc, "多层光学薄膜的逆向设计需要在离散材料、层数和厚度空间中寻找能够实现目标光谱的结构。对于斜入射条件，s 与 p 偏振的 Fresnel 响应和膜内干涉行为明显不同，使得分别设计两个膜系难以保证同一结构的联合性能。本文在 OptoGPT 生成式逆向设计框架上提出一种面向 60° 入射纯介质多层薄膜的联合偏振扩展模型。模型以 400–1100 nm、10 nm 间隔的联合光谱 [Rs, Ts, Rp, Tp] 作为 284 维条件，通过 s/p 双分支编码和融合模块驱动共享的自回归 Transformer 解码器，输出同一个材料—厚度 token 序列。为提高结构合法性和物理可信度，推理阶段使用纯介质 logits 掩码，并对多个候选结构进行精确传输矩阵法（TMM）重算与联合误差排序。我们构建了 500,000 个结构级无泄漏样本的数据集，并在独立测试样本上获得 100% 的合法解码率和 TMM 验证成功率，平均联合光谱误差为 0.03615。进一步的分布外压力测试显示，单候选、16 候选和 64 候选的平均联合 MAE 分别为 0.0930、0.0719 和 0.0634，说明物理重排序能够有效利用额外候选预算。有限厚度玻璃上的宽带高透射案例表明，p 偏振较易拟合，而 s 偏振仍是 60° 增透设计的主要瓶颈；同一模型在 70–80° 的性能也出现明显退化。本文因此建立了一个数据合同明确、生成结果可验证、并能诚实量化泛化边界的联合偏振薄膜逆向设计流程。")
para(doc, "关键词：多层薄膜；逆向设计；OptoGPT；联合 s/p 偏振；传输矩阵法；物理验证；分布外泛化")

para(doc, "1 引言", "Heading 1")
para(doc, "多层介质薄膜广泛用于增透涂层、滤光片、分布式布拉格反射镜、窄带选择器、成像系统和光伏器件等场景。给定期望的反射率或透射率光谱，设计者需要同时确定材料排列、层数和每层厚度。与正向模拟相比，逆向问题通常具有高维、离散、非凸和多解等特点：多个不同的膜层序列可能产生近似相同的光学响应，而局部优化方法又容易受初始结构和材料候选集影响。")
para(doc, "传输矩阵法（TMM）为平面分层结构提供了高效而精确的光学计算工具，但基于 TMM 的逐目标优化仍需要反复调用仿真器，难以在大规模目标空间中快速给出多种可制造候选。近年来，深度学习逆向设计通过学习目标光谱到结构序列的映射，能够显著降低单个目标的推理成本。OptoGPT 将材料和厚度组合编码为结构 token，并把多层膜序列化为自回归生成任务，从而支持可变层数、材料组合和多候选采样[1]。条件可逆神经网络等方法也表明，显式建模逆问题的多解性有助于生成结构集合，而不是只返回一个局部最优解[2]。")
para(doc, "然而，斜入射下的联合偏振设计仍存在三个关键困难。第一，s 和 p 偏振的界面反射和相位积累不同，分别预测两个结构并不能保证它们物理上是同一个膜系。第二，训练数据如果没有把同一结构的两种偏振响应沿特征维拼接，就无法形成严格的联合条件。第三，神经网络概率并不等价于光学性能，生成的 token 序列必须重新经过 TMM 检查，才能判断材料、层数、厚度和光谱响应是否同时满足要求。")
para(doc, "本文围绕上述问题，构建一个统一的联合 284D s+p OptoGPT 模型。主要贡献如下：")
bullets(doc, [
    "建立同一结构对应的联合偏振数据合同，以 [Rs, Ts, Rp, Tp] 组成 284 维条件，并通过结构哈希执行去重和确定性数据划分。",
    "设计 s/p 双光谱编码与融合模块，迁移预训练结构解码器，使模型输出一个共享的纯介质多层结构。",
    "在解码端使用纯介质 logits 掩码，并以精确 s+p TMM 对多候选结构进行物理重排序。",
    "同时报告同分布性能、严格分布外性能、多候选收益和有限玻璃应用结果，明确区分光谱等效解与原始结构恢复。",
])

para(doc, "2 物理问题与联合数据集", "Heading 1")
para(doc, "2.1 物理系统与光谱表示", "Heading 2")
para(doc, "考虑空气入射介质、纯介质多层膜和玻璃基底组成的平面结构。入射角固定为 60°，波长范围为 400–1100 nm，采样间隔为 10 nm，共 71 个波长点。允许材料包括 Al2O3、AlN、HfO2、MgF2、MgO、Si3N4、SiO2、Ta2O5、TiO2 和 ZnO；层数和厚度由结构 token 序列表示。对于任一结构，分别运行 s 偏振和 p 偏振 TMM，得到 Rs、Ts、Rp 和 Tp。联合条件为四条 71 点光谱沿特征轴拼接而成的 284 维向量。")
para(doc, "[Rs(λ1), …, Rs(λ71), Ts(λ1), …, Ts(λ71), Rp(λ1), …, Rp(λ71), Tp(λ1), …, Tp(λ71)] ∈ R^284。", align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "图 1 从物理应用、斜入射偏振分裂和模型流程三个层面概括了该问题。对于每个目标，模型只生成一个共享结构；s/p 两种响应在后续 TMM 中分别重算。")
figure(doc, "figure1.png", "图 1", "物理问题、联合 s+p OptoGPT 流程以及经典增透参考与 AI 候选的光谱对比。联合流程包括双偏振条件编码、纯介质候选生成和精确 TMM 重排序。")

para(doc, "2.2 数据生成、去重与划分", "Heading 2")
para(doc, "数据生成首先随机采样合法材料—厚度序列，再对同一物理结构执行两次偏振 TMM。四条 71 点光谱只在特征轴上拼接，不在样本轴上混合。每条结构记录同时保存 token 序列、材料和厚度元数据、入射角、波长网格以及联合光谱。结构级 SHA-256 哈希用于去重和划分，确保同一结构及其完全相同的光谱不会跨越训练集、开发集和测试集。")
table(doc, ["数据子集", "结构数", "光谱维度", "结构级泄漏"], [
    ["训练集", "400,006", "284", "0"],
    ["开发集", "50,104", "284", "0"],
    ["测试集", "49,890", "284", "0"],
    ["总计", "500,000", "284", "0"],
])
para(doc, "2.3 评价指标", "Heading 2")
para(doc, "对目标光谱 y 和候选结构经 TMM 重算得到的光谱 ŷ，分别定义 s/p 通道平均绝对误差 Es 和 Ep，并定义联合误差 Ejoint=(Es+Ep)/2。除误差外，本文还记录合法解码率、TMM 成功率、平均透射率、p05 透射率以及候选数量对应的 TMM 调用预算。MAE 衡量的是光谱匹配误差，不是结构 token 的恢复准确率。")

para(doc, "3 联合 284D OptoGPT 方法", "Heading 1")
para(doc, "3.1 结构序列化与共享解码器", "Heading 2")
para(doc, "每一层由材料和离散厚度组成一个 material_thickness token，多层结构按照沉积顺序序列化，并以 BOS 和 EOS 标记序列边界。共享的 decoder-only Transformer 根据联合光谱条件逐 token 生成结构。该表示避免固定层数输出的限制，使不同层数、材料顺序和厚度组合可以由同一生成器处理。")
para(doc, "3.2 s/p 双分支编码与参数迁移", "Heading 2")
para(doc, "联合模型将 284D 条件拆分为 s 分支 [Rs, Ts] 和 p 分支 [Rp, Tp]。两个分支分别通过光谱编码器提取特征，再由融合模块形成共享条件表示。预训练 OptoGPT 的 token embedding、六层自回归 decoder 和输出 generator 被迁移到联合模型；新增的 s/p 条件分支与融合参数进行初始化。训练分为融合预热和全量微调两个阶段，以降低新条件表示对已有结构语言表示的扰动。")
para(doc, "3.3 材料约束与物理重排序", "Heading 2")
para(doc, "纯介质约束在推理阶段通过 logits mask 实现：对金属和半导体 token 的 logits 设为负无穷，使最终采样不会输出禁用材料。对每个目标采样 N 个候选，依次进行序列合法性检查、结构去重和 s/p TMM 重算，然后按照 Ejoint 排序。该流程把神经网络作为高效候选提议器，把 TMM 作为物理裁判，避免用 token 概率代替真正的光学评价。")
para(doc, "3.4 训练设置与可复现性", "Heading 2")
para(doc, "正式训练使用 500,000 个联合样本、batch size 16、Adam 优化器、混合精度、label smoothing、ReduceLROnPlateau 学习率调度和差分学习率。最佳 checkpoint 由开发集 loss 选择，并保存模型、optimizer、scheduler、AMP scaler、epoch/step、随机数状态、架构版本、预训练权重哈希和数据 manifest。图 2 展示联合模型的训练阶段、材料词表约束和端到端物理评价合同。")
figure(doc, "figure2.png", "图 2", "联合 284D s+p 模型的训练进程、s/p 通道性能、纯介质词表约束及联合物理评价流程。图中四个性能柱均属于同一共享结构联合模型。")

para(doc, "4 训练收敛与联合模型验证", "Heading 1")
para(doc, "4.1 训练收敛", "Heading 2")
para(doc, "联合模型的 token-level train loss 从 4.6647 下降至 3.3164，dev loss 从 4.3498 下降至 3.0381。两个阶段均表现出持续下降趋势，最佳开发集 checkpoint 出现在第 10 个 epoch。需要区分 token prediction loss 与光谱 MAE：前者反映结构序列建模的优化过程，后者必须在结构生成后通过 TMM 重新计算。")
para(doc, "4.2 光谱重建与结构统计", "Heading 2")
para(doc, "图 3 给出联合模型在验证样本上的最佳和典型光谱重建。灰色曲线表示目标响应，彩色虚线表示由生成结构经 TMM 重新计算的 s/p 通道响应。200 个有效联合样本的平均 total MAE 为 0.03157，标准差为 0.02157。生成结构使用了十种允许的介质材料，层数范围为 1–15 层，平均层数为 6.50。")
figure(doc, "figure3.png", "图 3", "联合 284D 模型的代表性光谱重建、total-MAE 分布、材料使用、层数统计及 s/p 波长分辨误差。所有面板均来自同一共享结构联合模型。")
para(doc, "4.3 独立测试集上的联合性能", "Heading 2")
table(doc, ["指标", "结果", "解释"], [
    ["合法解码", "100/100", "抽样测试目标均生成合法结构"],
    ["TMM 验证", "100/100", "所有候选均完成精确 s/p 重算"],
    ["Es", "0.05084", "s 通道平均光谱误差"],
    ["Ep", "0.02145", "p 通道平均光谱误差"],
    ["Ejoint", "0.03615", "两通道误差的平均值"],
    ["平均 Ts", "0.50367", "随机测试目标的平均 s 透射率"],
    ["平均 Tp", "0.90179", "随机测试目标的平均 p 透射率"],
])
para(doc, "随机测试目标包含多种光谱形态，并不等价于统一的宽带高透射目标。因此平均 Ts 不应被解释为增透性能上限；它更直接地揭示了在联合条件下 s 通道仍然是主要误差来源。")

para(doc, "5 多候选、分布外泛化与结构多解性", "Heading 1")
para(doc, "5.1 候选预算的作用", "Heading 2")
para(doc, "为分离单次生成能力与物理搜索预算，本文在严格压力测试中分别比较 greedy 单候选、16 候选和 64 候选。每个候选均通过相同的 TMM 评价和去重流程，最终报告该目标候选池中的最小联合误差。")
table(doc, ["候选策略", "平均联合 MAE", "中位数 MAE", "最差 MAE"], [
    ["1 候选 greedy", "0.0930", "0.0889", "0.1722"],
    ["16 候选 + TMM 排序", "0.0719", "0.0714", "0.1300"],
    ["64 候选 + TMM 排序", "0.0634", "0.0643", "0.1131"],
])
para(doc, "从 1 个候选增加到 64 个候选，平均联合 MAE 从 0.0930 降至 0.0634。该改善不能简单归因于 Transformer 本身，而是生成多样性与精确物理筛选共同带来的系统收益。")
para(doc, "5.2 严格分布外测试", "Heading 2")
para(doc, "严格 OOD 集采用连续厚度、高层数、随机、渐变、强交替和双峰厚度结构，并在部分目标中加入小幅光谱噪声。该测试与可核验训练结构进行哈希去重，避免把训练结构的近似重现误认为泛化。64 候选结果中，s 偏振平均 MAE 为 0.0827，p 偏振平均 MAE 为 0.0440。强交替结构最难，表明模型对训练分布之外的快速相位变化和复杂层序列仍较敏感。")
para(doc, "5.3 光谱等效与结构非唯一性", "Heading 2")
para(doc, "在 60 个严格 OOD 目标上，生成结构完整恢复目标结构的比例为 0/60。这并不表示逆向设计失败，因为不同材料顺序和厚度组合可能产生近似光谱；它说明本文模型主要解决的是找到一个物理上等效的结构，而不是恢复数据生成时的原始结构。")

para(doc, "6 有限玻璃上的高透射设计与能力边界", "Heading 1")
para(doc, "为检验联合模型在更接近器件实际的基底条件下的表现，本文将候选结构放置于 500 μm 有限厚度玻璃上，并使用包含基底内部非相干反射的有限基底 TMM。目标设为 400–1100 nm 范围内 Rs=Rp=0.05、Ts=Tp=0.95 的平坦高透射响应。")
figure(doc, "figure4.png", "图 4", "有限厚度玻璃上的三个联合 s+p 高透射候选。每个候选均由同一平坦目标驱动并经有限玻璃 TMM 评价；最佳保留候选的平均 Ts=0.773、平均 Tp=0.982，尚未达到共同 0.95 目标。")
para(doc, "图 4 显示，p 偏振可以接近目标，而 s 偏振在有限玻璃条件下仍有明显反射损失。该结果应被理解为模型在工程约束下的能力边界，而不是已经实现宽带双偏振 95% 透射的证明。有限厚度基底引入的内部反射、玻璃—膜层界面以及 60° 下的偏振不对称共同增加了优化难度。")
para(doc, "进一步的角度扫描显示，60° 训练得到的结构在 70–80° 会明显退化；在代表性候选上，80° 附近 mean Ts 约为 0.3525、mean Tp 约为 0.5925。因而本文模型的适用范围应明确限定为 60° 条件，不能将单角度结果外推为 0–75° 宽角性能。")

para(doc, "7 讨论", "Heading 1")
para(doc, "本文的核心优势是把联合偏振问题写成一个数据合同清晰的共享结构生成任务，并把 TMM 放回推理闭环。与只报告神经网络 loss 的做法相比，结构合法性、TMM 成功率和候选预算都可以被独立审查。结果同时揭示了三个限制：")
bullets(doc, [
    "s 偏振在 60° 及更高角度下是主要瓶颈，说明当前十种致密介质和离散厚度空间未必包含足够宽的高角度增透解。",
    "多候选和 TMM 重排序会改善结果，但计算成本随候选数量增加；正式比较应将 TMM 调用次数作为统一预算。",
    "模型输出的是光谱等效结构而非唯一真实结构。面向制造时，还需要最小层厚、总厚度、材料兼容性和厚度扰动容差等约束。",
])
para(doc, "主动学习、宽角多角度条件模型、渐变或多孔折射率材料，以及制造误差闭环，是下一阶段值得开展的方向。当前主动学习基础设施已完成数据记录和去重模块，但尚未完成 ensemble、不确定性采集和闭环增广实验，因此不在本文中宣称主动学习性能提升。")

para(doc, "8 结论", "Heading 1")
para(doc, "本文提出并验证了一个面向 60° 斜入射纯介质多层薄膜的联合 s+p OptoGPT 模型。模型以同一结构产生的 284D [Rs, Ts, Rp, Tp] 光谱为条件，利用双偏振编码、共享自回归解码、纯介质 logits 掩码和 TMM 物理重排序生成可验证的多层结构。500,000 样本数据集的结构级划分、100% 的抽样解码与 TMM 成功率，以及严格 OOD 中随候选预算增加而下降的联合 MAE，共同构成了本文的主要证据链。")
para(doc, "结果同时表明，联合逆向设计并不等于所有偏振和角度都已解决：s 偏振是主要误差来源，有限玻璃高透射目标尚未完全达到，60° 模型在更大入射角下会退化。通过把这些边界纳入正文，本文给出了一个可复现、可物理核查且适合继续扩展的联合偏振薄膜逆向设计基线。")

para(doc, "参考文献", "Heading 1")
refs = [
    "Ma T, Wang H, Guo LJ. OptoGPT: A foundation model for inverse design in optical multilayer thin film structures. Opto-Electronic Advances. 2024;7(7):240062. doi:10.29026/oea.2024.240062.",
    "Luce A, Mahdavi A, Wankerl H, Marquardt F. Investigation of inverse design of multilayer thin-films with conditional invertible neural networks. Machine Learning: Science and Technology. 2023;4(1):015014. doi:10.1088/2632-2153/acb48d.",
    "Hong Y, Nicholls DP. Data-driven design of thin-film optical systems using deep active learning. Optics Express. 2022;30(13):22901. doi:10.1364/OE.459295.",
    "Khaireh-Walieh A, Langevin D, Bennet P, Teytaud O, Moreau A, Wiecha PR. A newcomer’s guide to deep learning for inverse design in nano-photonics. Nanophotonics. 2023;12(24):4387–4414. doi:10.1515/nanoph-2023-0527.",
    "Wiecha PR, Arbouet A, Girard C, Muskens OL. Deep learning in nano-photonics: inverse design and beyond. Photonics Research. 2021;9(5):B182–B200. doi:10.1364/PRJ.415960.",
    "Byrnes SJ. Multilayer optical calculations. arXiv:1603.02720. 2016. doi:10.48550/arXiv.1603.02720.",
    "Macleod HA. Thin-Film Optical Filters. 5th ed. Boca Raton: CRC Press; 2017.",
    "Goodfellow I, Bengio Y, Courville A. Deep Learning. Cambridge: MIT Press; 2016.",
    "Settles B. Active Learning Literature Survey. Madison: University of Wisconsin–Madison; 2009.",
]
for i, ref in enumerate(refs, 1):
    reference(doc, i, ref)

doc.save(OUT)
print(OUT)
