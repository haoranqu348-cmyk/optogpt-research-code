from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
SOURCE = Path(
    "/Users/quhaoran/Library/Containers/com.tencent.xinWeChat/Data/Documents/"
    "xwechat_files/wxid_5yofaowibzo22_2b5b/temp/drag/1.docx"
)
FORMULA_SOURCE = ROOT / "Physics_Validated_JP_OptoGPT_Manuscript_Body.docx"
OUTPUT = ROOT / "JP_OptoGPT_Concise_Teacher_Review.docx"


TEXT = {
    0: "TMM-Reranked Joint-Polarization Generative Inverse Design of Dielectric Multilayer Films",
    3: (
        "Oblique-incidence multilayer dielectric coatings exhibit strong polarization splitting between s- and "
        "p-polarized light. A practical inverse-design system must therefore generate one physical stack whose "
        "response is evaluated for both channels, rather than produce two incompatible polarization-specific "
        "coatings. We present JP-OptoGPT, a joint-polarization generative inverse-design framework adapted from "
        "OptoGPT and coupled to transfer-matrix-method (TMM) recomputation and physical reranking."
    ),
    4: (
        "JP-OptoGPT uses a dual-branch encoder to retain s/p channel identity before feature fusion, a deterministic "
        "logits mask that restricts output to ten dielectric materials, and multi-candidate generation followed by "
        "TMM evaluation. These components define the evaluated workflow; their individual contributions have not "
        "yet been isolated by matched ablation experiments."
    ),
    5: (
        "All simulations use 60° incidence and a 400–1100 nm grid with 10 nm spacing (71 wavelengths). The training "
        "corpus contains 500,000 legal dielectric stacks. On a fixed 100-target held-out sample, the best-of-up-to-32 "
        "procedure achieves mean joint spectral MAE Ejoint = 0.03615. In a combined out-of-distribution (OOD) stress "
        "test, changing from one greedy decode to a 64-candidate stochastic-and-ranking procedure changes mean joint "
        "MAE from 0.0930 to 0.0634; this difference reflects both decoding policy and candidate budget."
    ),
    6: (
        "For the finite-glass application, Candidate A raises mean s-polarized transmission from 0.6938 for bare "
        "glass to 0.7733 and raises the unpolarized mean from 0.8453 to 0.8776, while mean p-polarized transmission "
        "decreases. A subsequent double-sided optimization initialized from a generated candidate reaches an "
        "unpolarized mean of 0.9071. The role of model initialization is not isolated, and the target mean "
        "Ts = Tp = 0.95 remains unmet."
    ),
    7: (
        "The results establish an end-to-end computational proposal-and-ranking workflow for one-stack joint s/p "
        "design. They do not yet establish superiority to direct optimization, nearest-neighbor retrieval, or "
        "alternative model architectures under matched computational budgets."
    ),
    10: (
        "Multilayer optical coatings control reflection, transmission, and phase across broad spectra, but their "
        "design space grows rapidly with layer count and material combinations. Classical synthesis uses physical "
        "intuition together with numerical refinement or global optimization [2], while data-driven inverse design "
        "learns reusable mappings from target spectra to candidate structures [3]."
    ),
    11: (
        "OptoGPT represents each material-thickness pair as a structure token and formulates inverse design as "
        "conditional autoregressive generation [1]. At oblique incidence, however, s and p polarizations have "
        "different Fresnel coefficients and optical admittances. Designing the channels independently can therefore "
        "produce two incompatible coatings instead of one shared structure."
    ),
    12: "JP-OptoGPT addresses this shared-structure problem through three linked changes:",
    13: "Encode the paired 284-dimensional condition [Rs, Ts, Rp, Tp] with polarization-aware branches and feature fusion;",
    14: "Mask decoder logits to exclude prohibited metals and semiconductors, restricting generated layers to ten dielectrics;",
    15: "Generate multiple candidate stacks, recompute their s/p responses with TMM, and rank legal structures by joint spectral error.",
    16: (
        "Figure 1. Physical problem and JP-OptoGPT workflow. (a) Coated objective lens and magnified oblique-incidence "
        "interface showing the s/p field orientations. (b) Paired spectral input, polarization-aware encoding and "
        "fusion, autoregressive candidate generation, dielectric logits masking, and 60° TMM reranking. (c,d) "
        "Qualitative single-layer MgF2 responses at 0° and 60°, illustrating polarization splitting. (e) One "
        "model-generated high-transmission candidate at 60°. The MgF2 example is illustrative rather than a "
        "competitive quantitative baseline."
    ),
    19: (
        "The optical configuration is an isotropic multilayer between air and a 500 μm glass substrate. The fixed "
        "incidence angle is 60°, and the wavelength interval is 400–1100 nm with 10 nm spacing (Nλ = 71). The ten "
        "permitted dielectrics are Al2O3, AlN, HfO2, MgF2, MgO, Si3N4, SiO2, Ta2O5, TiO2, and ZnO."
    ),
    22: (
        "For lossless layers, Tq = 1 - Rq, so the four spectra are not independent and a compressed [Rs, Rp] "
        "condition contains the same optical information. The 284D form is retained as an explicit four-channel "
        "interface compatible with finite-substrate and future absorbing cases; its advantage over the reduced "
        "condition has not been established by a controlled ablation."
    ),
    24: "TMM computes the s- and p-polarized responses using polarization-dependent optical admittance [6]. Snell's law and the single-pass phase thickness are",
    26: "The normalized admittances are",
    28: (
        "The characteristic matrix, stack amplitude coefficients, and measurable intensity coefficients are given "
        "by Eqs. (5)–(7). The glass substrate is treated as incoherent. The same TMM implementation is used for label "
        "generation and candidate recomputation; no independent solver-to-solver numerical comparison is reported."
    ),
    36: (
        "For the finite-glass application, we additionally report wavelength-averaged transmission, the limiting "
        "polarization metric, and the equally weighted unpolarized mean:"
    ),
    37: "3. Physics-Constrained OptoGPT Extension",
    38: (
        "JP-OptoGPT inherits OptoGPT's material-thickness serialization and six-layer autoregressive decoder [1]. "
        "The present study evaluates the integrated system; it does not quantify the independent gain of each new "
        "component."
    ),
    40: (
        "The 284D input is divided into two 142D branches for [Rs, Ts] and [Rp, Tp]. Dedicated encoders preserve "
        "polarization identity before their features are fused to condition one shared decoder. Training uses a "
        "two-epoch fusion warm-up followed by full fine-tuning with differential learning rates. A direct "
        "single-branch concatenation encoder has not been tested under a matched protocol."
    ),
    42: (
        "The source vocabulary includes materials with appreciable extinction. At each material-generation step, "
        "the logits of prohibited materials are set to negative infinity before softmax, and probability is "
        "renormalized over the ten permitted dielectrics. The mask guarantees vocabulary legality but is not an "
        "accuracy ablation."
    ),
    44: (
        "Inference requests one greedy sequence followed by stochastic samples with top-k = 10, top-p = 0.9, and "
        "temperature = 1.0. EOS terminates decoding, and the maximum sequence length permits at most 20 layers. "
        "Invalid tokens and duplicate physical structures are removed before all surviving candidates are "
        "recomputed by TMM and ranked by Ejoint. Sequence likelihood is not used as the optical objective."
    ),
    45: (
        "Figure 2. Training and evaluation summary. (a) Archived model-development trajectory; the formal "
        "500,000-structure, ten-epoch run is a separate protocol, so the stages are not a controlled ablation. "
        "(b) Channel-separated reconstruction diagnostics. (c) Vocabulary restriction from 18 source materials to "
        "ten permitted dielectrics. (d) Paired-condition training and TMM-reranked evaluation workflow."
    ),
    48: (
        "The corpus contains 500,000 legal dielectric stacks with 1–20 layers. Materials are sampled from the "
        "ten-material vocabulary while preventing adjacent repeats, and thickness tokens span 10–300 nm in 10 nm "
        "increments. A SHA-256 hash of the physical structure produces 400,006 training, 50,104 development, and "
        "49,890 test structures, with no detected cross-split structure overlap."
    ),
    52: (
        "The OptoGPT checkpoint initializes training with random seed 42. Adam uses a base learning rate of 3×10^-5, "
        "label smoothing of 0.1, batch size 16, mixed precision, and ReduceLROnPlateau with factor 0.5 and minimum "
        "learning rate 1×10^-7. Training loss decreases from 4.6647 at epoch 1 to 3.3164 at epoch 10, while development "
        "loss reaches 3.0381 at epoch 10. Token cross-entropy is not optical accuracy; all reported spectra are "
        "evaluated after free-running decoding and TMM recomputation."
    ),
    55: (
        "Random seed 42 selects 100 targets without replacement from the held-out set. For each target, the inference "
        "procedure requests up to 32 candidates and retains the TMM-ranked structure with the lowest joint error. "
        "Every sampled target yields at least one legal TMM-evaluable stack. The reported errors are therefore "
        "best-of-up-to-32 system metrics, not one-shot decoder accuracy, and 100/100 validity does not imply a zero "
        "failure rate over the full test pool."
    ),
    56: "Table 2. Held-out performance on a fixed sample of 100 targets using best-of-up-to-32 TMM ranking.",
    59: (
        "A separate 200-sample auxiliary record gives mean total MAE = 0.03157 (standard deviation 0.02157), with "
        "generated layer counts from 1 to 15 (mean 6.50). This record is used only for descriptive statistics and is "
        "not pooled with the formal 100-target test because the two records do not share one declared evaluation protocol."
    ),
    60: (
        "Figure 3. Auxiliary spectral reconstruction and structural statistics for the joint 284D model. (a,b) "
        "Representative TMM-recomputed spectra. (c) Total-MAE distribution over 200 auxiliary samples. (d) Material "
        "frequency. (e) Generated layer-count distribution. (f) Wavelength-resolved s/p reconstruction error. The "
        "formal 100-target result in Table 2 remains the primary quantitative evaluation."
    ),
    63: (
        "The combined OOD set contains 60 targets with 15–20 layers, continuous off-grid thicknesses from 11 to "
        "499 nm, several structural families, and minor spectral perturbations. Because multiple shifts occur "
        "together, this is a combined stress test rather than a factorial analysis of individual OOD mechanisms. "
        "Structural non-overlap alone also does not quantify spectral novelty."
    ),
    65: (
        "Table 3 reports one greedy decode and two stochastic multi-candidate procedures. Mean joint MAE changes from "
        "0.0930 to 0.0719 and 0.0634 as the reported procedure changes. Because decoding policy and candidate budget "
        "change together, the 31.8% difference between the first and last rows cannot be attributed to candidate "
        "count alone. Under the 64-candidate procedure, mean Es = 0.0827 and mean Ep = 0.0440."
    ),
    69: (
        "Nineteen unique high-transmission candidates are collected. Sixteen enter the dense 1° scan from 0° to "
        "80°; the remaining three are not part of the full-range gate denominator. None of the 16 densely scanned "
        "candidates passes the combined full-range thresholds. One representative stack has mean transmission near "
        "0.937 for both polarizations near normal incidence, but at 80° its values fall to Ts ≈ 0.3525 and "
        "Tp ≈ 0.5925. This diagnostic bounds the tested candidates rather than all possible model outputs."
    ),
    71: (
        "The application target is flat over 400–1100 nm with Rs = Rp = 0.05 and Ts = Tp = 0.95 on a 500 μm "
        "incoherent glass substrate. From 4096 requested sequences, 2032 unique legal structures pass TMM evaluation, "
        "and three representative candidates are shown. Bare glass is the physical reference, but it is not a "
        "competitive coating baseline; no single-layer MgF2 or conventionally optimized multilayer is included."
    ),
    74: (
        "Candidate A improves the limiting s-polarized and unpolarized mean transmissions while slightly reducing the "
        "already high p-polarized mean. A subsequent differential-evolution optimization uses a generated candidate "
        "as the initialization for front and rear coatings and reaches an unpolarized mean of 0.9071. This refined "
        "coating is a post-hoc engineering result rather than a direct JP-OptoGPT output. Because random, quarter-wave, "
        "and multi-restart initializations are not compared under a common budget, the benefit of model seeding is not established."
    ),
    75: (
        "Figure 4. Finite-glass application and double-sided refinement. The upper panels show three candidates selected "
        "from 2032 unique TMM-valid structures. The lower panels compare bare glass, Candidate A, and the refined "
        "double-sided coating. Mean Ts increases from 0.6938 to 0.7733 and 0.8350; the unpolarized mean increases from "
        "0.8453 to 0.8776 and 0.9071; mean Tp decreases from 0.9967 to 0.9818 and 0.9792. The target mean "
        "Ts = Tp = 0.95 remains unmet."
    ),
    77: (
        "Contribution relative to OptoGPT. JP-OptoGPT adapts the variable-length generator [1] to a shared-structure "
        "s/p task through paired conditioning, polarization-aware fusion, deterministic material constraints, and "
        "multi-candidate TMM reranking. The 284D representation is a uniform interface rather than four independent "
        "degrees of information in the lossless limit. Direct-concatenation, reduced-condition, no-warm-up, and "
        "random-initialization controls are not yet available."
    ),
    78: (
        "Comparison boundary. Reinforcement learning and probabilistic inverse models provide alternative ways to "
        "construct or represent multilayer solutions [4,5]. The present work does not include matched-compute "
        "comparisons with those methods, nearest-neighbor retrieval, or conventional thin-film optimization."
    ),
    79: (
        "Data efficiency. Uniform sampling of the 500,000-structure corpus provides broad coverage but is inefficient "
        "for the narrow high-angle, high-transmission region. Active learning is a plausible extension [7], although "
        "no active-learning gain is demonstrated here."
    ),
    80: (
        "Physical and fabrication limits. TMM assumes planar, homogeneous, isotropic films and fixed complex n-k "
        "tables. It does not model roughness, interdiffusion, or process-dependent optical constants. The discrete "
        "thickness grammar does not represent continuous fabrication errors, and the ten-material vocabulary lacks "
        "the ultra-low-index options often used for wide-angle antireflection [8]. Solver-to-solver agreement and "
        "experimental measurements remain separate validation requirements."
    ),
    82: (
        "JP-OptoGPT adapts OptoGPT to one-stack joint s/p inverse design using paired conditioning, dielectric-constrained "
        "generation, and multi-candidate TMM reranking. On a fixed 100-target held-out sample, the best-of-up-to-32 "
        "procedure produces at least one legal TMM-evaluable candidate for every target and reaches mean "
        "Ejoint = 0.03615."
    ),
    83: (
        "In the combined OOD stress test, changing from one greedy decode to a 64-candidate stochastic-and-ranking "
        "procedure changes mean joint MAE from 0.0930 to 0.0634. In the finite-glass case, Candidate A improves the "
        "limiting-polarization and unpolarized metrics relative to bare glass while reducing mean Tp; subsequent "
        "double-sided optimization reaches an unpolarized mean of 0.9071 but does not attain the equal 0.95 target."
    ),
    84: (
        "The work therefore establishes computational feasibility rather than comparative superiority. The main "
        "remaining evidence gaps are module ablations, equal-budget traditional-optimization and retrieval baselines, "
        "larger statistical evaluation, controlled optimizer-initialization tests, independent solver comparison, "
        "fabrication robustness, and experimental validation."
    ),
}


REFERENCES = [
    "[1] T. Ma, H. Wang, and L. J. Guo, \"OptoGPT: a foundation model for inverse design in optical multilayer thin film structures,\" Opto-Electron. Adv. 7, 240062 (2024), doi:10.29026/oea.2024.240062.",
    "[2] H. A. Macleod, Thin-Film Optical Filters, 5th ed. (CRC Press, 2017).",
    "[3] S. So, T. Badloe, J. Noh, J. Rho, and J. Bravo-Abad, \"Deep learning enabled inverse design in nanophotonics,\" Nanophotonics 9, 1041–1057 (2020), doi:10.1515/nanoph-2019-0474.",
    "[4] H. Wang, Z. Zheng, C. Ji, and L. J. Guo, \"Automated multi-layer optical design via deep reinforcement learning,\" Mach. Learn.: Sci. Technol. 2, 025013 (2021), doi:10.1088/2632-2153/abc327.",
    "[5] R. Unni, K. Yao, and Y. Zheng, \"Deep convolutional mixture density network for inverse design of layered photonic structures,\" ACS Photonics 7, 2703–2712 (2020), doi:10.1021/acsphotonics.0c00630.",
    "[6] A. Luce, A. Mahdavi, F. Marquardt, and H. Wankerl, \"TMM-Fast, a transfer matrix computation package for multilayer thin-film optimization: tutorial,\" J. Opt. Soc. Am. A 39, 1007–1013 (2022), doi:10.1364/JOSAA.450928.",
    "[7] Y. Hong and D. P. Nicholls, \"Data-driven design of thin-film optical systems using deep active learning,\" Opt. Express 30, 22901–22916 (2022), doi:10.1364/OE.459295.",
    "[8] K. Pfeiffer, L. Ghazaryan, U. Schulz, and A. Szeghalmi, \"Wide-angle broadband antireflection coatings prepared by atomic layer deposition,\" ACS Appl. Mater. Interfaces 11, 21887–21894 (2019), doi:10.1021/acsami.9b03125.",
]


FORMULA_INDEX = {1: 13, 2: 17, 3: 19, 4: 21, 5: 23, 6: 25, 7: 27, 8: 31, 9: 33, 10: 35, 11: 38, 12: 40}


def replace_paragraph_text(paragraph, text):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.3)


def replace_figure_caption(paragraph, text):
    drawings = [deepcopy(node) for node in paragraph._p.xpath(".//w:drawing")]
    paragraph.clear()
    for drawing in drawings:
        paragraph.add_run()._r.append(drawing)
    paragraph.add_run("\n" + text)


def copy_formula_into(target, source):
    ppr = target._p.pPr
    for child in list(target._p):
        if child is not ppr:
            target._p.remove(child)
    for child in source._p:
        if child.tag != qn("w:pPr"):
            target._p.append(deepcopy(child))


def insert_formula_after(paragraph, source):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    copy_formula_into(result, source)
    return result


def insert_body_row(table, position, values):
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        cell.text = value
    table._tbl.remove(row._tr)
    table._tbl.insert(position + 1, row._tr)


def no_split_rows(table):
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if row_index == 0 and tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.3)
    normal.paragraph_format.space_after = Pt(2.3)
    normal.paragraph_format.line_spacing = 1.02

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if i == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(5)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(15)
                run.font.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)
            continue
        if text == "Authors:":
            paragraph.paragraph_format.space_after = Pt(5)
            for run in paragraph.runs:
                run.font.bold = True
            continue
        if text in {"Abstract", "References"} or re.match(r"^\d+\.\s", text):
            level = 1 if re.match(r"^\d+\.\s", text) or text in {"Abstract", "References"} else 2
            paragraph.style = doc.styles["Heading 1"]
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(6 if level == 1 else 4)
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)
            continue
        if re.match(r"^\d+\.\d+\s", text):
            paragraph.style = doc.styles["Heading 2"]
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(1.5)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10.3)
                run.font.bold = True
                run.font.color.rgb = RGBColor(46, 116, 181)
            continue
        if text.startswith("Figure "):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(4)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(7.6)
            continue
        if text.startswith("Table "):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(1.5)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(8.3)
                run.font.bold = True
            continue
        if text.startswith("[") and re.match(r"^\[\d+\]", text):
            paragraph.paragraph_format.space_after = Pt(0.4)
            paragraph.paragraph_format.line_spacing = 0.92
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(7.1)
            continue
        for run in paragraph.runs:
            if not run.font.name:
                run.font.name = "Arial"
            if not run.font.size:
                run.font.size = Pt(9.3)

    for table in doc.tables:
        no_split_rows(table)
        for r, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(7.5)
                        if r == 0:
                            run.font.bold = True


def main():
    doc = Document(SOURCE)
    formula_doc = Document(FORMULA_SOURCE)
    original_paragraphs = list(doc.paragraphs)

    for index, text in TEXT.items():
        if index in {16, 45, 60, 75}:
            replace_figure_caption(original_paragraphs[index], text)
        else:
            replace_paragraph_text(original_paragraphs[index], text)

    for i, ref in enumerate(REFERENCES, start=86):
        replace_paragraph_text(original_paragraphs[i], ref)

    formula_targets = {1: 21, 2: 25, 4: 27, 8: 31, 9: 33, 10: 35}
    for number, target_index in formula_targets.items():
        copy_formula_into(original_paragraphs[target_index], formula_doc.paragraphs[FORMULA_INDEX[number]])

    after_eq2 = insert_formula_after(original_paragraphs[25], formula_doc.paragraphs[FORMULA_INDEX[3]])
    cursor = original_paragraphs[28]
    for number in (5, 6, 7):
        cursor = insert_formula_after(cursor, formula_doc.paragraphs[FORMULA_INDEX[number]])
    cursor = original_paragraphs[36]
    for number in (11, 12):
        cursor = insert_formula_after(cursor, formula_doc.paragraphs[FORMULA_INDEX[number]])

    # Make the abbreviated held-out table explicit about validity and inference policy.
    table2 = doc.tables[1]
    insert_body_row(table2, 0, ["Inference policy", "Best of up to 32", "1 greedy + stochastic samples + TMM ranking"])
    insert_body_row(table2, 0, ["Valid/TMM-evaluable targets", "100/100", "Fixed sampled targets only"])

    style_document(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
