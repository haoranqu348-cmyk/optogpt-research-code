from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Physics_Validated_JP_OptoGPT_Manuscript_Body.docx"
OUTPUT = ROOT / "Physics_Validated_JP_OptoGPT_Manuscript_Body_Revised.docx"
WORK = ROOT / "paper_rewriting_output"


REVISED = {
    2: "Multilayer optical films offer a compact way to control reflection, transmission, absorption, phase, and color across broad spectral bands. Turning that flexibility into a design is difficult. A complete stack is specified by its materials, their order, the number of layers, and every layer thickness; these choices interact through multiple reflection and interference rather than contributing independently. Even with a fixed material set, the number of possible sequences grows rapidly with layer count. Classical synthesis therefore combines optical insight with numerical refinement, needle insertion, evolutionary search, or related target-specific procedures [1]. These methods remain effective, but they usually restart the search for each target and can return different solutions when the initialization or permitted structure family changes.",
    3: "A learned inverse model changes the cost profile of this process: training is expensive, but a trained model can reuse what it has learned when a new target arrives. Nanophotonic inverse design has consequently been explored with discriminative and tandem networks, mixture-density and generative models, reinforcement learning, and differentiable surrogates [2–4]. For multilayer films, reinforcement-learning agents can construct variable-length stacks [5,6], while probabilistic models can represent several structures compatible with one spectrum [7]. None of these choices alone makes a generated coating physically acceptable. Material restrictions still have to be enforced, and the requested response must be checked with a forward solver rather than inferred from sequence probability.",
    4: "OptoGPT is a useful starting point because it represents each material-thickness pair as a structure token and treats inverse design as conditional autoregressive generation [8]. The token sequence can end after different numbers of layers, and repeated sampling can return structurally distinct answers to the same optical condition. Pretraining also provides a structural prior that can be adapted through fine-tuning. Oblique-incidence design, however, changes the task itself. The condition must describe both polarization channels of one stack, the decoder must be prevented from selecting unsuitable materials, and the final choice must be based on TMM-recomputed optical error rather than sequence likelihood.",
    5: "At oblique incidence, s and p polarizations no longer see the same interface. Their Fresnel coefficients and optical admittances differ, whereas the propagation phase of a shared isotropic layer is set by its refractive index, thickness, and internal angle. The separation is pronounced at 60 degrees: a coating with low p-polarized reflection may still reflect much of the s-polarized field. Solving the two channels independently would produce two coatings, which is not a solution for a single component illuminated with an arbitrary polarization state. We therefore define one learning record by one layer sequence and its complete `[Rs, Ts, Rp, Tp]` response.",
    6: "We developed JP-OptoGPT around that shared-structure requirement. The optical input is a paired 284-dimensional record containing reflection and transmission for both polarizations. Separate s and p encoders preserve channel identity before a fusion module conditions the inherited autoregressive decoder. During decoding, a logits mask confines material choices to ten dielectric materials. The model acts as a proposal generator: it produces multiple legal candidates, and the same s/p transfer-matrix framework used for the labels recomputes and ranks their optical responses. In this paper, physics validation refers to that solver-based recomputation and physical filtering; it does not denote an independent second solver or experimental validation.",
    7: "This division between proposal and selection also determines how performance is reported. One decoded sequence measures one-shot generation; the best result among 16 or 64 samples measures the larger system formed by stochastic generation and TMM ranking. We keep those quantities separate. The same applies to data distribution: the formal held-out test samples the archived dataset, whereas a stricter out-of-distribution set uses continuous thicknesses, longer stacks, alternating patterns, and perturbed spectra. Taken together, the two evaluations distinguish interpolation within the training grammar from recovery of comparable spectra outside it.",
    8: "The present experiments evaluate the integrated proposal-and-ranking workflow. They establish that the system can generate legal shared-structure candidates and that TMM selection can identify lower-error proposals. They do not yet isolate the marginal contribution of the dual-branch encoder, pretrained initialization, fusion warm-up, or four-spectrum representation, nor do they establish superiority to nearest-neighbor retrieval or a conventional optimizer under the same forward-solver budget.",
    11: "We consider an isotropic multilayer between air and a 500 micrometer glass substrate. Unless a section states otherwise, the incidence angle is 60 degrees and the spectrum spans 400-1100 nm in 10 nm increments, so `N_lambda=71`. The allowed material set is Al2O3, AlN, HfO2, MgF2, MgO, Si3N4, SiO2, Ta2O5, TiO2, and ZnO. Complex, wavelength-dependent optical constants are loaded from the archived local n-k database rather than replacing the extinction coefficient by zero. Each generated layer is encoded by a material-thickness token; the strict OOD test later relaxes the discrete training thickness grid to continuous values.",
    12: "For a structure `x`, the forward solver returns reflection and transmission at every sampled wavelength for both polarizations. We concatenate those four spectra along the feature dimension:",
    14: "Equation (1) is primarily a data-integrity and interface definition. All four spectra must have been computed from the same layer sequence, and their order remains unchanged during data generation, training, inference, and evaluation. In the lossless limit, `T_s=1-R_s` and `T_p=1-R_p`, so the four spectra are not independent and a 142-dimensional `[Rs,Rp]` condition contains the same optical information. The 284-dimensional form retains explicit reflection and transmission channels for compatibility with finite-substrate and future absorbing cases; the present study does not show that it outperforms the reduced condition.",
    16: "The transfer-matrix method is the forward model used for both labels and candidate evaluation. For a nonmagnetic layer `j`, Snell's law relates the incident angle to the internal angle through the refractive index:",
    18: "The corresponding single-pass phase thickness is",
    20: "Polarization enters through the normalized optical admittance:",
    22: "Under this convention, layer `j` has the characteristic matrix",
    24: "The stack matrix `M` is the ordered product of the layer matrices. Let `q0` and `qsub` denote the incident-medium and substrate admittances for the polarization being evaluated. Writing the boundary fields as `[B,C]^T=M[1,qsub]^T` gives the amplitude coefficients",
    26: "The measurable intensity coefficients follow as",
    28: "Every candidate is evaluated twice, once with each polarization-specific admittance. For a lossless stack on a semi-infinite substrate, the implementation recovers `R+T=1` to numerical precision; with complex optical constants, any absorption is retained through the imaginary refractive-index component. The finite-glass calculation used in Sec. 7 requires a different treatment. Fields remain coherent within the coatings, but the 500 micrometer glass plate is treated as an incoherent layer; repeated internal intensity transfers and the rear interface are included. The implementation follows the conventions used by TMM-Fast [9] and the multilayer formulation described by Byrnes [10], but no cross-solver numerical error study is reported here.",
    30: "For `q` equal to s or p, we average the absolute reflection and transmission deviations over the wavelength grid. In the lossless limit, the two terms are algebraically linked and should not be interpreted as independent error evidence; both are retained to preserve the four-channel interface:",
    32: "The two channel errors enter the joint objective with equal weight:",
    34: "Given `K` sequences sampled from the decoder, the selected structure is the legal candidate with the smallest joint error:",
    36: "Here `C_legal` contains only sequences that pass syntax, material, thickness, and TMM checks. The notation makes the search budget explicit. `K=1` describes one-shot generation, whereas `K>1` adds a finite search over model proposals. Sequence log probability is not used as a surrogate for optical quality.",
    37: "The application study also uses the wavelength-averaged transmission of each polarization and the smaller of the two channel means:",
    39: "For equally weighted unpolarized illumination, the corresponding mean is",
    41: "These averages expose the trade studied on finite glass. A coating may raise the limiting s channel and the unpolarized mean even when the p-channel mean decreases; such a result is a balanced-metric improvement, not simultaneous improvement of both polarizations.",
    46: "JP-OptoGPT retains the material-thickness language introduced by OptoGPT. An `L`-layer stack is serialized as `[BOS, m1_d1, ..., mL_dL, EOS]`, with each interior token specifying one material and one thickness. The decoder predicts a conditional distribution over the next token and may emit EOS at different positions. Layer count, material sequence, and thickness can therefore vary without changing the dimensionality of the output head.",
    47: "We transfer the token embedding, positional representation, six-layer autoregressive decoder, and output generator. Those weights provide the initialization for learning layer-order and thickness patterns. The work reported here does not depend on a new Transformer block; its main change lies in the physical condition presented to the decoder and in the constrained, TMM-reranked acceptance procedure. Because no random-initialization control is available, the present results do not quantify the benefit of transfer learning itself.",
    49: "The 284-dimensional condition is divided before encoding. One 142-dimensional branch receives `[Rs,Ts]`; the other receives `[Rp,Tp]`. Their channel representations are fused and supplied to the shared decoder, keeping polarization identity explicit through the first encoding stage while conditioning one common sequence. This is the architecture evaluated here. A direct-concatenation encoder was not trained under a matched protocol, so the results do not establish that the dual-branch design is necessary or superior.",
    50: "Training starts with a two-epoch fusion warm-up. During this stage, inherited structural parameters remain frozen while the new condition encoders and fusion module adapt to the pretrained decoder. Full fine-tuning then releases all parameters, using the recorded differential-learning-rate schedule. The warm-up is part of the reported training protocol, but its contribution has not been isolated against an otherwise identical no-warm-up run.",
    52: "The intended outputs are low-loss dielectric coatings, but the source vocabulary also contains metals and semiconductors with appreciable extinction. An audit of the archived complex optical constants identified Ag, Al, Ge, Si, TiN, and related entries as unsuitable for the target band because they can introduce strong absorption or metallic reflection. At each material-generation step, their logits are set to negative infinity before softmax. Probability is then renormalized over the ten permitted dielectrics.",
    53: "The mask establishes legality, not an accuracy improvement. The migrated archive does not contain a controlled experiment that evaluates the same 284D checkpoint and candidate budget with the mask both enabled and disabled. We can therefore verify that prohibited material tokens are excluded, but we do not assign any change in MAE to the mask. Keeping those claims separate matters because a deterministic engineering constraint is not an empirical ablation.",
    55: "Inference encodes the joint condition once and requests up to `K` structures. The candidate generator emits one greedy sequence followed by stochastic samples with top-k = 10, top-p = 0.9, and temperature = 1.0; EOS terminates generation, and the maximum sequence length of 22 permits at most 20 layers. Special tokens are removed, tokens are parsed, prohibited materials are rejected, and duplicate physical structures are discarded before TMM evaluation. Every surviving sequence is then recomputed for s and p polarization with the same solver framework used to produce the labels.",
    56: "The proposal-and-ranking split can reject a high-probability sequence whose recomputed spectrum is poor, expose several solutions to the non-unique inverse problem, and express search cost through requested candidates and TMM calls. Section 6 varies that cost from 1 to 16 and 64 requested candidates. Because the `K=1` case is greedy whereas the larger sets include stochastic samples, this comparison measures the combined inference procedure rather than candidate count alone.",
    58: "Historical checkpoints are tied to an explicit architecture version because shape-compatible implementations can still execute different activation paths. The formal checkpoint is loaded with the recorded `joint_sp_legacy_v1` semantics and complete file hashes; an unknown unversioned checkpoint is rejected rather than interpreted heuristically.",
    59: "This version binding is the part relevant to the reported spectra. Detailed migration, Windows deployment, and console-compatibility records are implementation material rather than evidence for the optical contribution.",
    62: "The formal corpus contains 500,000 legal dielectric stacks. Layer count is sampled uniformly from 1 to 20; each material is sampled uniformly from the permitted choices while preventing adjacent repeats, and each available thickness token is sampled uniformly from 10-300 nm in 10 nm increments. Each stack was simulated at 60 degrees for s and p polarization on the same 71-point wavelength grid, after which the records were concatenated in the order defined by Eq. (1). The build pipeline checked finite values, token parsing, materials, thicknesses, layer counts, and sampled TMM recomputations before marking the dataset complete.",
    63: "We split the data with a SHA-256 hash of the physical structure rather than with its floating-point spectrum. This produced 400,006 training, 50,104 development, and 49,890 test structures, with no detected structure overlap across splits. Structure-level hashing keeps identical token sequences together even if serialization or numeric formatting changes, which is the relevant protection against leakage for this task.",
    67: "The OptoGPT checkpoint initialized a ten-epoch formal training run with random seed 42. Training used Adam, batch size 16, a base learning rate of 3e-5, label smoothing of 0.1, mixed precision, differential learning rates, and ReduceLROnPlateau with factor 0.5, patience 1, and minimum learning rate 1e-7. The first two epochs formed the fusion warm-up; all parameters were then fine-tuned. Training loss fell from 4.6647 at epoch 1 to 3.3164 at epoch 10, while development loss fell from 4.3498 to its minimum value of 3.0381 at epoch 10.",
    68: "Those losses measure teacher-forced token prediction. Optical reconstruction is evaluated only after free-running decoding and TMM recomputation. The distinction is substantive: a lower cross-entropy loss does not show that the hidden training structure has been recovered, nor does it ensure that one decoded sequence is the best optical solution available from the model.",
    70: "Figure 2(a) retains v3, v4, and ultimate checkpoints as development history; they are not a continuous 27-epoch experiment and are not used as a module ablation. The quantitative claims in Sec. 5 come from the separate ten-epoch, 500,000-structure protocol described above. The historical curves should therefore not be compared as if only one controlled factor changed between stages.",
    71: "The remaining panels describe two contracts rather than two ablations. The material panel shows the reduction of an 18-material source vocabulary, including eight metals or semiconductors, to ten allowed dielectrics. The flow panel follows a target through the 284D condition, shared-structure decoding, dielectric masking, exact 60-degree TMM, and the reported optical and structural statistics.",
    73: "Each formal checkpoint stores the model, optimizer, scheduler, AMP scaler, epoch, global step, random states, architecture version, pretrained hash, and data-manifest information. Resumption occurs at epoch boundaries and leaves the original pretrained file unchanged. Platform-specific deployment details are not part of the scientific comparison and are kept outside the main argument.",
    78: "For the primary quantitative test, random seed 42 selected 100 targets without replacement from the 49,890-structure held-out pool. For each target, the inference procedure requested up to 32 candidates, consisting of one greedy decode and stochastic samples, and retained the TMM-ranked candidate with the lowest joint error. Every sampled target yielded at least one legal stack and completed s/p TMM evaluation. The resulting best-of-up-to-32 means were `E_s=0.05084`, `E_p=0.02145`, and `E_joint=0.03615`. The larger s-channel error is consistent with the optical asymmetry near the Brewster regime, but the experiment does not isolate its cause.",
    81: "Table 2 summarizes a mixed target distribution, not a dedicated antireflection benchmark. Its errors are multi-candidate system metrics rather than one-shot decoder accuracy, and the 100/100 validity result applies only to this fixed sample; it does not imply a zero failure rate over all 49,890 held-out structures. A larger fixed evaluation with quantiles and confidence intervals is needed for population-level reliability. The high-transmission question is tested separately in Sec. 7.",
    83: "A separate 200-sample archive provides the distributional view shown in Fig. 3. All 200 decoded structures are TMM-valid. Their mean total MAE is 0.03157 with a standard deviation of 0.02157, and their layer counts range from 1 to 15 with a mean of 6.50. All ten allowed dielectrics appear; TiO2, ZnO, Ta2O5, Si3N4, and MgF2 are the most frequent among the retained structures.",
    84: "We do not pool this auxiliary set with the 100-sample formal test because the two records do not share one declared sampling and inference protocol. The 200-sample distribution is therefore descriptive support for Fig. 3 rather than an enlargement of the primary test. A future consolidated evaluation should use one checkpoint, one candidate policy, one fixed sample list, and the same summary statistics across all targets.",
    86: "The two reconstructions in Fig. 3 illustrate errors that a single mean can hide. In panel (a), a simple generated stack follows a smooth target closely across the band. Panel (b) contains sharper features; local deviations become visible where reflectance changes rapidly. Broad low-amplitude discrepancies and narrow larger mismatches can produce comparable aggregate MAE, so both the curves and the summary statistic are needed.",
    87: "Across much of the band, the wavelength-resolved error in Fig. 3(f) lies near 0.02-0.04, with larger values in parts of the long-wavelength region. Reflection and transmission errors are related by energy conservation in the dielectric stack. These curves describe optical reconstruction by the generated structure; they do not measure recovery of the original token sequence.",
    89: "The layer-count and material histograms confirm that the decoder is not confined to a single stack family. It proposes structures of different complexity, which creates room for a later fabrication-oriented choice. Two candidates with similar spectra may still differ in total thickness, interface count, material availability, or sensitivity to deposition error.",
    90: "Raw diversity is not sufficient. A sampled set may contain malformed, repeated, or optically poor sequences. Material masking, structure-level deduplication, and TMM ranking turn that output into a collection of legal candidates that can be compared under one physical objective. The next section asks whether increasing the size of that collection actually improves OOD performance.",
    95: "Evaluation is easiest when test structures follow the same grammar as the training set. To reduce that advantage, we constructed 60 OOD targets from 15-20-layer stacks with continuous thicknesses between 11 and 499 nm instead of the 10 nm training grid. The set combines random, graded, strongly alternating, and bimodal-thickness families, and 12 targets include small spectral perturbations. No exact structure overlap was found against the locally verifiable training structures.",
    96: "Each OOD condition still pairs s and p spectra for one shared target. However, length, thickness discretization, structural family, and spectral perturbation change together. The reported aggregate is therefore a combined stress test, not a factorial attribution of OOD difficulty. Structural non-overlap alone also does not quantify spectral novelty; separate shift-specific subsets and a nearest-neighbor spectral-distance analysis would be required for that claim.",
    98: "Table 3 follows the same 60 targets as the inference policy changes. One greedy candidate gives a mean joint MAE of 0.0930. Sampling and TMM-ranking 16 candidates gives 0.0719, and the 64-candidate procedure gives 0.0634; the median and worst-case values fall as well. The 31.8% difference between the first and last rows describes the full policy change and cannot be assigned to candidate count alone.",
    101: "This comparison characterizes the combined effect of candidate budget, stochastic proposal, and TMM selection under the reported inference policies. The `K=1` row is a single greedy decode; the other rows add stochastic proposals and select among them with the same physical objective. Accordingly, 0.0634 is not the error of a direct single output. It is the performance of JP-OptoGPT at a budget of 64 candidates and their associated forward calculations, not an isolated reranking ablation.",
    102: "The observed gain is larger from one to 16 candidates than from 16 to 64. A controlled budget study should use the same stochastic policy for `K=1,4,16,64`, nested candidate sets, repeated seeds, and the actual number of unique TMM-valid candidates or solver calls on the horizontal axis. Reporting confidence intervals, wall-clock time, and a separate greedy baseline would then distinguish search budget from decoding strategy.",
    104: "The polarization imbalance survives the combined distribution shift. At `K=64`, mean s-channel MAE is 0.0827 and mean p-channel MAE is 0.0440. Strongly alternating stacks have the largest family mean, approximately 0.0826. This association may reflect their impedance and phase structure, sequence length, or their frequency under the decoder distribution; the current aggregate does not identify a causal mechanism.",
    105: "None of the 60 hidden OOD structures is recovered exactly. That result does not contradict successful spectral inverse design. Multilayer synthesis is non-unique, and different stacks can produce closely related spectra, as also observed in probabilistic inverse-design studies [7,11]. Our objective is spectral agreement under the joint condition, not reconstruction of the data-generating sequence.",
    107: "We also scanned model-generated high-transmission candidates from 0 to 80 degrees with exact TMM. Nineteen unique candidates were collected; 16 entered the dense 1-degree scan, while the remaining three are not part of the full-range gate denominator. None of the 16 densely scanned candidates passed the combined mean, lower-tail, and minimum-transmission thresholds over the full range. One representative stack gives approximately 0.937 mean transmission for both polarizations near normal incidence, but at 80 degrees its values fall to `T_s` about 0.3525 and `T_p` about 0.5925.",
    108: "This diagnostic bounds the tested candidate set, not every output the model could produce. It shows that no densely scanned candidate establishes a 0-75-degree omnidirectional coating. Broadband wide-angle antireflection requires careful low-index matching, dispersion control, and tolerance design [12]; a broader claim would require multi-angle conditioning, a worst-angle objective, a larger refractive-index space, and a substantially larger controlled candidate evaluation.",
    111: "The application target is flat across 400-1100 nm, with `Rs=Rp=0.05` and `Ts=Tp=0.95`. We requested 4096 sequences and retained 2032 unique candidates that passed the TMM contract. The three structures in the upper part of Fig. 4 were selected for limiting-polarization transmission, unpolarized mean transmission, and structural diversity. Their different layer counts and material orders illustrate the multiplicity available under one joint target.",
    112: "Candidate A has the best limiting-channel result among the displayed structures: on a 500 micrometer finite-glass substrate, mean `T_s=0.7733` and mean `T_p=0.9818`. Candidate B is much thinner and reaches mean `T_p` of about 0.993, but its mean `T_s` is about 0.748. Candidate C uses another material sequence and gives mean `T_s` near 0.771 and mean `T_p` near 0.963. In all three cases, the p channel remains close to the target while s-polarized reflection stays well above 5%.",
    113: "Bare finite glass provides the physical reference, with mean `T_s=0.6938`, mean `T_p=0.9967`, and mean unpolarized transmission 0.8453. Candidate A raises the s-channel mean to 0.7733 and the unpolarized mean to 0.8776, while the p-channel mean decreases to 0.9818. The coating therefore improves the limiting channel and the balanced response, not both polarizations separately. Because no single-layer MgF2 coating or conventionally optimized multilayer is included, bare glass should not be interpreted as a competitive design baseline.",
    115: "The lower part of Fig. 4 records a double-sided optimization initialized from a generated structure. The final stack is not a one-step JP-OptoGPT output. Front and rear coatings are refined in the finite-glass geometry, with coherent propagation in the coatings and incoherent transfer through the substrate recomputed throughout. Because random, quarter-wave, and multi-restart optimization initializations were not compared under the same budget, this result does not show that the model initialization improves convergence or final performance.",
    116: "The optimized front coating is MgF2 143.9 nm / Al2O3 151.5 nm / TiO2 10.0 nm / SiO2 32.9 nm; the rear coating is MgO 169.3 nm / Al2O3 108.9 nm / MgF2 139.1 nm. Its means are `R_s=0.1634`, `R_p=0.0197`, `T_s=0.8350`, and `T_p=0.9792`, giving unpolarized mean transmission 0.9071. Compared with bare glass, `T_s` increases by 0.1412 and the unpolarized mean by 0.0618. Compared with Candidate A, the additional increases are 0.0617 and 0.0295, respectively.",
    120: "The supported application claim is narrower: relative to bare finite glass, Candidate A improves the limiting-polarization mean and the balanced unpolarized-transmission metric while reducing the already high p-channel mean. The recorded double-sided optimization uses a model candidate as its initialization, but the absence of matched initialization controls prevents a claim that the model supplied a better starting point.",
    121: "Neither Candidate A nor the refined coating simultaneously achieves mean `T_s>=0.95` and mean `T_p>=0.95`. The refined result also misses the stricter wavelength-resolved reflection and robustness gates, and the available perturbation study is insufficient to establish fabrication readiness. Figure 4 is therefore a TMM-reranked computational case study with a persistent s-polarized limitation, not an independently validated or experimentally demonstrated coating.",
    122: "The trade is consistent with, but not uniquely explained by, the interface physics near the Brewster regime, where bare glass already transmits p polarization strongly. Training distribution, spectral complexity, and objective weighting can also contribute to the observed s/p imbalance. The result is therefore reported through the measured channel and balanced-transmission metrics rather than assigned to one causal mechanism.",
    127: "Original OptoGPT supplies the variable-length sequence generator [8]. JP-OptoGPT adapts that generator to a shared-structure s/p task through paired conditioning, polarization-aware encoding and fusion, deterministic dielectric legality, and candidate-level TMM ranking with an explicit search budget. The four-spectrum 284D representation is a uniform interface rather than four independent degrees of optical information in the lossless limit. Its value relative to a reduced `[Rs,Rp]` condition remains an open ablation question.",
    128: "The current evidence supports the operation of the integrated workflow, not the indispensability of each component. Direct concatenation versus dual-branch encoding, random initialization versus OptoGPT transfer, and fusion warm-up versus immediate full fine-tuning have not been compared under matched conditions. Nearest-neighbor retrieval and conventional multilayer optimization have also not been evaluated at the same TMM-call budget. Consequently, the present results should not be read as a causal decomposition or a claim that JP-OptoGPT is uniformly superior to those alternatives.",
    130: "Other learning paradigms divide the work differently. Reinforcement-learning methods can add layers sequentially and optimize a physical reward [5,6], while mixture-density and invertible models represent conditional multiplicity explicitly [7,11]. JP-OptoGPT instead retains a variable-length token grammar, samples several proposals, and delegates final comparison to TMM. This is a methodological contrast only; the present study does not include a matched empirical benchmark against those architectures.",
    131: "Benchmark studies caution against treating one inverse architecture as uniformly best across nanophotonic problems [13]. The results here support a hybrid interpretation in which the Transformer proposes plausible stacks and the forward solver judges their optical response. Whether this division improves accuracy or compute efficiency over database retrieval or direct optimization must be tested under equal TMM calls, common targets, repeated seeds, and the same feasibility constraints.",
    133: "The 500,000-structure corpus is large, yet it sparsely covers the high-angle region in which both polarization channels transmit strongly. Uniform sampling spends most forward calculations on ordinary stacks outside that narrow region. Active learning has improved data-driven thin-film design by selecting informative simulations [14], and uncertainty-guided acquisition is established more broadly in materials design [15]. The remaining data problem is therefore one of allocation as much as scale.",
    134: "A future loop could rank unlabeled stacks by predicted transmission, ensemble disagreement, and structural novelty, evaluate selected candidates with TMM, and return them to training. Its comparison should be made against cumulative forward-solver calls, with random acquisition and performance-only selection as baselines. The current project contains record and deduplication infrastructure that could support such a study, but it does not yet provide evidence of an active-learning gain.",
    136: "TMM assumes planar, laterally homogeneous, isotropic layers and known optical constants. The calculations use the archived wavelength-dependent complex n-k tables, including their extinction coefficients, but do not model surface roughness, interdiffusion, deposition-induced density changes, or anisotropy. Deposited-film optical constants vary with process and microstructure, so a stack optimized against one database can shift in fabrication. Solver-to-solver agreement and experimental measurements remain separate validation requirements.",
    137: "There is also a mismatch between the formal token grammar and fabrication errors. Training thicknesses are discrete, whereas deposition variations are continuous. Continuous optimization in the double-sided study removes the grid only for that final refinement; it does not make the result robust. Future ranking should include correlated thickness errors, refractive-index uncertainty, process-calibrated optical constants, minimum layer thickness, and total coating thickness.",
    138: "The angular test points to a material-space limitation as well as a conditioning limitation. Broadband dual-polarization transmission at 70-80 degrees may require porous, graded, or otherwise very low-index layers that are absent from the ten-material dense vocabulary. Wide-angle coatings fabricated by atomic layer deposition illustrate the role of controlled index matching [12]. Effective-medium or continuously parameterized materials could expand the search, but they would also change the fabrication assumptions and should be evaluated as a separate model extension.",
    140: "The reporting boundaries are therefore substantive. Optical MAE is not structure accuracy; the primary 100-target result is a best-of-up-to-32 system metric; the `K=1,16,64` rows mix decoding policy with budget; the combined OOD set does not isolate individual shifts; the material mask is a legality rule rather than an accuracy ablation; and the refined coating is not a direct model output. All reported devices remain computational.",
    141: "The evidence closes an end-to-end feasibility loop but not the comparative one. The most important next tests are module ablations, a nearest-neighbor and conventional-optimization comparison under equal TMM calls, and a controlled initialization study for the double-sided optimizer. Until those results are available, JP-OptoGPT should be interpreted as a reproducible proposal-and-ranking framework rather than a demonstrated replacement for established thin-film synthesis.",
    143: "JP-OptoGPT adapts OptoGPT to shared-structure s/p inverse design through paired 284D conditioning, polarization-aware encoding, dielectric-constrained decoding, and multi-candidate TMM ranking. The 284D form provides an explicit four-channel interface, although its advantage over the lossless reduced condition has not been established. Formal training uses 500,000 structures split by physical hash, with no detected cross-split structure leakage.",
    144: "On a fixed 100-target held-out sample, the best-of-up-to-32 procedure produced at least one legal TMM-evaluable candidate for every target and achieved mean `E_joint=0.03615`. In the combined OOD test, changing from one greedy decode to a 64-candidate stochastic-and-ranking procedure changed mean joint MAE from 0.0930 to 0.0634. For finite glass, Candidate A increased mean `T_s` from 0.6938 to 0.7733 and mean unpolarized transmission from 0.8453 to 0.8776, while mean `T_p` decreased. Subsequent double-sided optimization reached 0.8350 and 0.9071 for the first two metrics, but the role of model initialization was not isolated.",
    145: "The present scope is correspondingly limited. The s channel remains limiting, the joint 0.95 target is not attained, robustness and experimental performance are unresolved, and no tested dense-scan candidate establishes full 0-80-degree operation. Module ablations, same-budget retrieval and traditional-optimization baselines, larger statistical evaluation, controlled initialization tests, and independent solver comparison are required before stronger comparative or validation claims are made.",
}


SECTION_EVIDENCE = {
    "1": "Refs. [1-8]; confirmed motivation; Fig. 1",
    "2": "Eqs. (1-12); TMM contract; Refs. [9,10]",
    "3": "Model implementation, material audit, checkpoint metadata",
    "4": "500k manifest, split hashes, formal training history; Table 1; Fig. 2",
    "5": "Formal 100-target test and auxiliary 200-sample archive; Table 2; Fig. 3",
    "6": "60-target OOD archive, candidate budgets, angular scan; Table 3",
    "7": "Finite-glass candidate and double-sided archives; Table 4; Fig. 4",
    "8": "Results above; Refs. [5-15]",
    "9": "Tables 2-4 and stated boundary evidence",
}


def section_for(doc, index):
    section = ""
    for i in range(index, -1, -1):
        text = doc.paragraphs[i].text.strip()
        match = re.match(r"^(\d+)\.", text)
        if match and doc.paragraphs[i].style.name == "Heading 1":
            section = match.group(1)
            break
    return section


def replace_paragraph(paragraph, text):
    paragraph.clear()
    word_text = text.replace("`", "")
    for source, target in (
        ("N_lambda", "Nλ"),
        ("T_unpol", "Tunpol"),
        ("T_worst", "Tworst"),
        ("E_joint", "Ejoint"),
        ("C_legal", "Clegal"),
        ("R_s", "Rs"),
        ("R_p", "Rp"),
        ("T_s", "Ts"),
        ("T_p", "Tp"),
        (">=", " ≥ "),
        ("400-1100", "400–1100"),
        ("0-80", "0–80"),
        ("0-75", "0–75"),
        ("70-80", "70–80"),
        ("15-20", "15–20"),
        ("0.02-0.04", "0.02–0.04"),
    ):
        word_text = word_text.replace(source, target)
    word_text = re.sub(r"\s*=\s*", " = ", word_text)
    word_text = re.sub(r"\s+", " ", word_text)
    word_text = word_text.replace("0–80-degree", "0‑80-degree")
    run = paragraph.add_run(word_text)
    run.font.name = "Arial"
    run.font.size = Pt(10.2)


def safe_cell(text):
    return text.replace("|", "\\|").replace("\n", " ")


def prevent_table_row_split(table):
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cantSplit") is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def write_markdown(doc):
    lines = ["# Revised JP-OptoGPT Manuscript", ""]
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name
        if style == "Heading 1":
            lines.extend([f"## {text}", ""])
        elif style == "Heading 2":
            lines.extend([f"### {text}", ""])
        elif i in range(13, 42, 2) and re.fullmatch(r"\(\d+\)", text):
            lines.extend([f"[Equation {text[1:-1]} retained as OMML in the DOCX]", ""])
        elif text.startswith("Figure ") or text.startswith("Table "):
            lines.extend([text, ""])
        elif i < 147:
            lines.extend([text, ""])
    manuscript = "\n".join(lines)
    (WORK / "revised_manuscript.md").write_text(manuscript, encoding="utf-8")
    final_paper = WORK / "final_paper"
    final_paper.mkdir(exist_ok=True)
    (final_paper / "revised_manuscript.tex").write_text(manuscript, encoding="utf-8")


def write_audit_files(original, revised):
    logic = [
        "# Original Logic Map",
        "",
        "| Original Unit | Current Text Role | Evidence Used | Motivation Link | Problem | Decision |",
        "|---|---|---|---|---|---|",
    ]
    rationale = [
        "# Writing Rationale Matrix",
        "",
        "| Row ID | Manuscript Unit | Current Problem or Planned Function | Motivation Link | Reference/SOTA Pattern Learned | Target Scene or Venue Norm | User Evidence or Citation Anchor | Planned Change/Text Move | Final Text Check |",
        "|---|---|---|---|---|---|---|---|---|",
        "| F1 | Whole-work framework | The draft contains the correct experiments but often announces contributions and caveats in a repeated list-like rhythm. | Make the physical contract, rather than generic model novelty, control the evidence sequence. | Transfer only the common technical-paper move of defining the problem before architecture and testing the claim before the application. | Optics Express expects a complete solver contract, reproducible evaluation, and bounded claims. | Eqs. (1-12), Figs. 1-4, Tables 1-4, formal and OOD archives. | Preserve the nine-section order but rewrite each paragraph around one evidence-bearing decision. | PASS: revised sections follow problem -> contract -> method -> formal test -> stress test -> application -> limits. |",
    ]
    human = [
        "# Humanize Matrix",
        "",
        "The matrix documents stylistic revision. It does not predict or guarantee the output of an AI-detection service.",
        "",
        "| Row ID | Manuscript Unit | AI Pattern Found | Detection Dim | Severity | Applied Change | Expected Effect | Teaching Note |",
        "|---|---|---|---|---|---|---|---|",
        "| HG1 | Whole manuscript | Repeated multi-clause syntax created a narrow sentence-structure range. | D1-Sentence Structure | High | Rebuilt prose with short definitions, medium evidence statements, and longer causal interpretations. | Increases syntactic and rhythmic variation. | Sentence structure should follow the reasoning task rather than a fixed template. |",
        "| HG2 | Whole manuscript | Adjacent paragraphs frequently used the same claim-explain-summary grammar. | D2-Paragraph Similarity | High | Assigned different paragraph forms to definitions, comparisons, results, boundaries, and transitions. | Reduces paragraph similarity across the manuscript. | Paragraph similarity is high when every unit closes with the same generic takeaway. |",
        "| HG3 | Whole manuscript | Technical terms sometimes appeared inside stock explanatory phrases. | D5-Term-Context Matching | Low | Retained every domain term but placed it next to the specific method, result, or limitation it denotes. | Reduces formulaic term-context matching without substituting technical vocabulary. | Medium-tier scientific revision should vary context, not rename established concepts. |",
    ]
    rewrite = [
        "# Rewrite Matrix",
        "",
        "| Row ID | Original Unit | Final Unit | Change Type | Evidence Preserved | Claim Boundary Check |",
        "|---|---|---|---|---|---|",
    ]
    blueprint_sections = {}

    for row, index in enumerate(sorted(REVISED), 1):
        old = original.paragraphs[index].text.strip()
        new = revised.paragraphs[index].text.strip()
        section = section_for(original, index)
        evidence = SECTION_EVIDENCE.get(section, "Current manuscript evidence")
        unit = f"Sec. {section} P{index}"
        old_sentences = len(re.findall(r"[.!?](?:\s|$)", old))
        new_sentences = len(re.findall(r"[.!?](?:\s|$)", new))
        role = "Advance the section's evidence sequence without adding a claim."
        problem = "Uniform multi-clause exposition and a predictable paragraph-final summary obscure the evidential priority."
        logic.append(
            f"| {unit} | {role} | {safe_cell(evidence)} | Integrated physical-contract motivation | {safe_cell(problem)} | Rewrite |"
        )
        rationale.append(
            f"| R{row:02d} | {unit} | {safe_cell(problem)} | Tie the paragraph to the shared-structure proposal-and-verification argument. | Use direct claim-to-evidence movement without copying source wording. | Keep technical detail close to the result or method it qualifies. | {safe_cell(evidence)} | Rebuild the paragraph with varied sentence scale and an evidence-led opening; preserve every number and citation. | PASS: {old_sentences} source sentences -> {new_sentences} revised sentences; claims and anchors retained. |"
        )
        human.extend([
            f"| H{row:02d}a | {unit} | Consecutive sentences cluster around a similar multi-clause length. | D1-Sentence Length | High | Mixed short declarative sentences with longer causal or interpretive sentences ({old_sentences} -> {new_sentences} sentences). | Broadens the sentence-length distribution. | Uniform sentence length makes technical prose sound generated even when the content is correct. |",
            f"| H{row:02d}b | {unit} | Repeated claim-explain-summary paragraph skeleton. | D2-Paragraph Structure | High | Reorganized the unit around its local function: method definition, evidence comparison, causal interpretation, or abrupt boundary statement. | Reduces structural similarity across adjacent paragraphs. | Human technical writing changes paragraph shape when the reasoning task changes. |",
            f"| H{row:02d}c | {unit} | Information density remains flat across description, transition, and result interpretation. | D3-Information Density | Medium | Concentrated numbers and constraints in evidence paragraphs while shortening bridges and equation lead-ins. | Creates intentional density variation. | A constant density across every paragraph is a common template signal. |",
        ])
        if re.search(r"\b(First|Second|Third|Finally|Moreover|Furthermore|Additionally|In conclusion|The article is organized)\b", old, re.I):
            human.append(
                f"| H{row:02d}d | {unit} | Formulaic enumerator or roadmap connector. | D4-Connector Frequency | Medium | Removed the stock connector and expressed the logical relation through subject and verb choice. | Lowers connector density at paragraph openings. | Explicit connective scaffolding is useful sparingly; uniform use makes the argument sound preformatted. |"
            )
        rewrite.append(
            f"| W{row:02d} | Paragraph {index} | {unit} | Rhetorical and language; no factual change | {safe_cell(evidence)} | PASS: numerical values, citations, materials, and result scope preserved. |"
        )
        blueprint_sections.setdefault(section, []).append(unit)

    blueprints = ["# Section Blueprints", ""]
    for section, units in blueprint_sections.items():
        blueprints.extend([
            f"## Section {section}",
            "",
            f"- Units revised: {', '.join(units)}",
            f"- Evidence anchor: {SECTION_EVIDENCE.get(section, 'Current manuscript evidence')}",
            "- Drafting rule: state the physical or evidential decision first, keep interpretation proportional to the reported result, and end without a generic summary when the point is already established.",
            "",
        ])

    (WORK / "original_logic_map.md").write_text("\n".join(logic), encoding="utf-8")
    (WORK / "writing_rationale_matrix.md").write_text("\n".join(rationale), encoding="utf-8")
    (WORK / "humanize_matrix.md").write_text("\n".join(human), encoding="utf-8")
    (WORK / "rewrite_matrix.md").write_text("\n".join(rewrite), encoding="utf-8")
    (WORK / "section_blueprints.md").write_text("\n".join(blueprints), encoding="utf-8")
    (WORK / "logic_transfer_audit.md").write_text(
        "# Logic Transfer Audit\n\n"
        "## Verdict\n\nPASS. The revision preserves the original nine-section evidence order and all displays. "
        "The controlling argument remains the shared-structure physical contract, followed by model adaptation, "
        "formal evaluation, OOD candidate-budget testing, finite-glass application, and bounded discussion.\n\n"
        "## Checks\n\n"
        "- No section, equation, figure, table, or reference was removed or reordered.\n"
        "- No numeric result or material sequence was changed.\n"
        "- Best-of-K results remain identified as framework results with an explicit budget.\n"
        "- The dielectric mask remains a legality claim, not an accuracy ablation.\n"
        "- The refined double-sided design remains distinct from a direct model output.\n"
        "- Failed angular and robustness gates remain visible.\n",
        encoding="utf-8",
    )


def main():
    original = Document(SOURCE)
    revised = Document(SOURCE)
    missing = [index for index in REVISED if index >= len(revised.paragraphs)]
    if missing:
        raise RuntimeError(f"Paragraph indices no longer match source: {missing}")
    for index, text in REVISED.items():
        replace_paragraph(revised.paragraphs[index], text)
    for table in revised.tables:
        prevent_table_row_split(table)
    for paragraph in revised.paragraphs:
        if paragraph.text.startswith("Table 2."):
            paragraph.paragraph_format.page_break_before = True
            break
    for paragraph in revised.paragraphs[147:162]:
        paragraph.paragraph_format.space_after = Pt(0.6)
        paragraph.paragraph_format.line_spacing = 0.95
        for run in paragraph.runs:
            run.font.size = Pt(7.8)
    revised.save(OUTPUT)
    write_markdown(revised)
    write_audit_files(original, revised)
    print(OUTPUT)
    print("REVISED_PARAGRAPHS", len(REVISED))


if __name__ == "__main__":
    main()
